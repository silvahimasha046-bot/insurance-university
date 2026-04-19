#!/usr/bin/env python3
"""
6BUIS020C Final Project Report Generator
Insurance University - AI-Powered Insurance Recommendation Platform
Student: Himaha Silva (IIT: 20212128 / UoW: 18712342/1)
Generates a fully formatted Microsoft Word (.docx) report.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Constants ───────────────────────────────────────────────────────────────
OUTPUT_FILE = "Insurance_University_Final_Report_Himaha_Silva.docx"
DIAGRAM_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "report_diagrams")
STUDENT_NAME = "Himaha Silva"
IIT_ID = "20212128"
UOW_ID = "18712342/1"
SUBMISSION_DATE = "22nd April 2026"
FONT_NAME = "Times New Roman"
FONT_SIZE_NORMAL = Pt(12)
FONT_SIZE_H1 = Pt(16)
FONT_SIZE_H2 = Pt(14)
FONT_SIZE_H3 = Pt(13)
LINE_SPACING = 1.5


# ─── Helper Class ────────────────────────────────────────────────────────────
class ReportBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_margins()

    def _setup_margins(self):
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _setup_styles(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = FONT_NAME
        font.size = FONT_SIZE_NORMAL
        pf = style.paragraph_format
        pf.line_spacing = LINE_SPACING
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(6)

        for level, size in [("Heading 1", FONT_SIZE_H1), ("Heading 2", FONT_SIZE_H2), ("Heading 3", FONT_SIZE_H3)]:
            s = self.doc.styles[level]
            s.font.name = FONT_NAME
            s.font.size = size
            s.font.bold = True
            s.font.color.rgb = RGBColor(0, 0, 0)
            s.paragraph_format.space_before = Pt(12)
            s.paragraph_format.space_after = Pt(6)
            s.paragraph_format.line_spacing = LINE_SPACING

    def add_page_break(self):
        self.doc.add_page_break()

    def h1(self, text):
        return self.doc.add_heading(text, level=1)

    def h2(self, text):
        return self.doc.add_heading(text, level=2)

    def h3(self, text):
        return self.doc.add_heading(text, level=3)

    def p(self, text, bold=False, italic=False, alignment=None, font_size=None):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = font_size or FONT_SIZE_NORMAL
        run.bold = bold
        run.italic = italic
        para.paragraph_format.line_spacing = LINE_SPACING
        para.paragraph_format.alignment = alignment or WD_ALIGN_PARAGRAPH.JUSTIFY
        return para

    def centered(self, text, bold=False, font_size=None):
        return self.p(text, bold=bold, font_size=font_size, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def bullet(self, text):
        para = self.doc.add_paragraph(style="List Bullet")
        para.clear()
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_NORMAL
        para.paragraph_format.line_spacing = LINE_SPACING
        return para

    def numbered(self, text):
        para = self.doc.add_paragraph(style="List Number")
        para.clear()
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_NORMAL
        para.paragraph_format.line_spacing = LINE_SPACING
        return para

    def table(self, headers, rows):
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Light Grid Accent 1"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header row
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.bold = True
                    run.font.name = FONT_NAME
                    run.font.size = Pt(11)
        # Data rows
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = tbl.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val)
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.font.name = FONT_NAME
                        run.font.size = Pt(11)
        return tbl

    def placeholder_image(self, caption):
        self.p(f"[INSERT IMAGE: {caption}]", italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.p(f"Figure: {caption}", italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    def insert_image(self, image_path, caption, width=Inches(6)):
        """Insert an actual image with a caption below it."""
        if os.path.exists(image_path):
            self.doc.add_picture(image_path, width=width)
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.p(caption, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            self.placeholder_image(caption)

    def save(self, path):
        self.doc.save(path)
        print(f"Report saved to: {path}")


# ─── Report Content ──────────────────────────────────────────────────────────
def build_report():
    rb = ReportBuilder()

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════
    rb.p("")
    rb.p("")
    rb.p("")
    rb.centered("6BUIS020C \u2013 Final Project Report", bold=True, font_size=Pt(14))
    rb.p("")
    rb.p("")
    rb.centered("Insurance University: An AI-Powered Insurance\nRecommendation Platform", bold=True, font_size=Pt(22))
    rb.p("")
    rb.p("")
    rb.centered(f"Student: {STUDENT_NAME} (IIT ID: {IIT_ID} / UoW ID: {UOW_ID})", font_size=Pt(14))
    rb.p("")
    rb.p("")
    rb.centered(
        "This report is submitted in partial fulfilment of the requirements for the\n"
        "BSc (Hons) Business Information Systems",
        font_size=Pt(12),
    )
    rb.p("")
    rb.p("")
    rb.centered("Business School", bold=True, font_size=Pt(14))
    rb.p("")
    rb.centered("Informatics Institute of Technology, Sri Lanka\nin collaboration with\nUniversity of Westminster, UK", font_size=Pt(13))
    rb.p("")
    rb.p("")
    rb.centered(SUBMISSION_DATE, font_size=Pt(14))

    # ═══════════════════════════════════════════════════════════════════════
    # DECLARATION
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("Declaration")
    rb.p(
        "This report has been prepared based on my own work. Where other published and unpublished "
        "source materials have been used, these have been acknowledged in references."
    )
    rb.p("")
    rb.p(f"Word Count: [To be updated after final review]")
    rb.p(f"Student Name: {STUDENT_NAME}")
    rb.p(f"Date of Submission: {SUBMISSION_DATE}")

    # ═══════════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("Abstract")
    rb.p(
        "The insurance industry in Sri Lanka is characterised by low product penetration, manual advisory processes, "
        "and a lack of transparency in product suitability assessment. Customers face difficulty understanding which "
        "insurance plans align with their financial profile, family obligations, and risk appetite. This project, "
        "Insurance University, addresses these challenges by delivering an AI-powered, web-based insurance "
        "recommendation platform that automates the end-to-end customer journey from data capture through to "
        "personalised plan recommendations and Know-Your-Customer (KYC) document submission."
    )
    rb.p(
        "The system employs a four-tier microservices architecture consisting of an Angular 21 single-page application, "
        "a Spring Boot 3.3.5 RESTful API backend with MySQL 8.4 persistence (16 JPA entities), and a FastAPI-based AI engine "
        "with an integrated agentic chat subsystem. "
        "The AI engine implements a multi-stage scoring pipeline incorporating Classification and Regression Tree (CART) "
        "underwriting rules, collaborative filtering heuristics, and trainable feature weights derived from historical "
        "policy outcome data. A novel agentic chat subsystem powered by a Groq-hosted LLM (Meta LLaMA-4 Scout) provides "
        "natural-language insurance advisory through five specialised tools (knowledge base, database query, web search, "
        "calculator, and policy scoring), with FAISS-based vector memory for long-term conversational context. "
        "Server-Sent Events (SSE) enable real-time streamed responses. "
        "The platform supports two user roles\u2014Customer and Administrator\u2014with JWT-based "
        "authentication and role-based access control. Key results demonstrate the system\u2019s ability to rank products "
        "with explainable scoring, provide real-time premium estimates with affordability analysis, deliver conversational "
        "insurance guidance via an agentic AI chat, and offer an "
        "administrative console for product management, eligibility rule configuration, and AI model lifecycle governance. "
        "The project successfully delivers a functional prototype that streamlines insurance advisory services and "
        "lays the groundwork for data-driven product innovation in the Sri Lankan insurance market."
    )
    rb.insert_image(os.path.join(DIAGRAM_DIR, "01_architecture.png"), "Figure 1: System Architecture Diagram")

    # ═══════════════════════════════════════════════════════════════════════
    # ACKNOWLEDGEMENTS
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("Acknowledgements")
    rb.p(
        "I would like to express my sincere gratitude to my project supervisor at the Informatics Institute of "
        "Technology for their invaluable guidance, constructive feedback, and continuous encouragement throughout "
        "the development of this project. Their expertise in software engineering and information systems was "
        "instrumental in shaping the direction and quality of this work."
    )
    rb.p(
        "I am also thankful to the academic staff at IIT and the University of Westminster for providing a "
        "rigorous curriculum that equipped me with the knowledge and skills necessary to undertake a project of "
        "this scope and complexity."
    )
    rb.p(
        "Special thanks go to my family and friends for their unwavering support and patience during the "
        "demanding phases of development and documentation. Their encouragement kept me motivated to deliver "
        "the best possible outcome."
    )
    rb.p(
        "Finally, I acknowledge the open-source communities behind Angular, Spring Boot, FastAPI, and the "
        "numerous libraries used in this project, whose contributions made rapid development possible."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("Table of Contents")
    rb.p("[Right-click and select \u2018Update Field\u2019 in Microsoft Word to auto-generate the Table of Contents]", italic=True)
    # Insert a TOC field code
    paragraph = rb.doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    run._r.append(fldChar1)
    run2 = paragraph.add_run()
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>'.format(nsdecls('w')))
    run2._r.append(instrText)
    run3 = paragraph.add_run()
    fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))
    run3._r.append(fldChar2)
    run4 = paragraph.add_run("Press F9 or right-click \u2192 Update Field to populate this Table of Contents")
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run5 = paragraph.add_run()
    fldChar3 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run5._r.append(fldChar3)

    # ═══════════════════════════════════════════════════════════════════════
    # LIST OF FIGURES
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("List of Figures")
    figures = [
        ("Figure 1", "System Architecture Diagram"),
        ("Figure 2", "Fishbone (Ishikawa) Diagram \u2013 Insurance Advisory Problems"),
        ("Figure 3", "Onion Model \u2013 Stakeholder View"),
        ("Figure 4", "High-Level Architecture Diagram"),
        ("Figure 5", "Use Case Diagram (21 Use Cases)"),
        ("Figure 6", "Class Diagram \u2013 Backend Entity Model (16 Entities)"),
        ("Figure 7", "Sequence Diagram \u2013 Get Insurance Recommendations"),
        ("Figure 8", "Entity-Relationship Diagram (16 Tables)"),
        ("Figure 9", "Wizard Step 1 \u2013 Personal Information"),
        ("Figure 10", "Wizard Step 2 \u2013 Financial Information"),
        ("Figure 11", "Wizard Step 3 \u2013 Insurance Preferences"),
        ("Figure 12", "Recommendations View"),
        ("Figure 13", "Product Comparison View"),
        ("Figure 14", "Proposal Upload View"),
        ("Figure 15", "Admin Dashboard"),
        ("Figure 16", "Admin Training Management"),
        ("Figure 17", "Admin Logs Viewer"),
        ("Figure 18", "Open Chat (Agentic AI) Interface"),
    ]
    rb.table(["Figure No.", "Title"], [[f[0], f[1]] for f in figures])

    # ═══════════════════════════════════════════════════════════════════════
    # LIST OF TABLES
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("List of Tables")
    tables_list = [
        ("Table 1", "Technology Stack Summary"),
        ("Table 2", "Functional Requirements"),
        ("Table 3", "Non-Functional Requirements"),
        ("Table 4", "JPA Entity Summary"),
        ("Table 5", "AI Scoring Weight Parameters"),
        ("Table 6", "Persistent vs Ephemeral Data Classification"),
        ("Table 7", "Test Plan \u2013 Features to Be Tested"),
        ("Table 8", "Test Cases"),
        ("Table 9", "Glossary of Terms"),
    ]
    rb.table(["Table No.", "Title"], [[t[0], t[1]] for t in tables_list])

    # ═══════════════════════════════════════════════════════════════════════
    # PART I: PROJECT DESCRIPTION
    # ═══════════════════════════════════════════════════════════════════════

    # ── SECTION 1 ─────────────────────────────────────────────────────────
    rb.add_page_break()
    rb.h1("1  Project Overview")
    rb.p(
        "Insurance University is a web-based, AI-driven insurance recommendation platform designed to modernise "
        "the way individuals in Sri Lanka discover, evaluate, and apply for insurance products. The platform "
        "replaces the traditional broker-mediated advisory process with a guided digital experience that captures "
        "a customer\u2019s personal, financial, and lifestyle information through a structured multi-step wizard, "
        "then leverages an artificial intelligence scoring engine to produce personalised, ranked product "
        "recommendations with transparent explanations."
    )
    rb.p(
        "The system is architected as a set of loosely coupled microservices: an Angular single-page application "
        "provides the user interface for both customer and administrator journeys, including an agentic open-chat "
        "interface with real-time SSE streaming; a Spring Boot RESTful API "
        "manages authentication, session orchestration, chat services, product catalogues, eligibility rules, and audit logging; "
        "a FastAPI-based AI engine executes a multi-stage scoring pipeline that includes CART-based underwriting "
        "checks, collaborative filtering heuristics, premium estimation, affordability analysis, and lapse-risk "
        "prediction, and additionally hosts an agentic chat subsystem powered by a Groq-hosted LLM with five "
        "specialised tools and FAISS-based long-term vector memory; and a MySQL relational database (16 tables) "
        "persists all transactional and reference data including chat messages."
    )
    rb.p(
        "Beyond the customer-facing recommendation journey, Insurance University provides a comprehensive "
        "administration console that enables insurance product managers to maintain the product catalogue, "
        "define eligibility rules, upload historical policy outcome datasets for AI model training, promote "
        "trained models into production, monitor system-wide event logs, and track unmatched customer needs "
        "to inform future product development. The platform thus serves as both an operational tool for "
        "insurance distribution and a strategic analytics asset for product innovation."
    )

    # ── SECTION 2 ─────────────────────────────────────────────────────────
    rb.h1("2  The Purpose of the Project")

    rb.h2("2a  The User Business or Background of the Project Effort")
    rb.p(
        "The insurance industry in Sri Lanka, regulated by the Insurance Regulatory Commission of Sri Lanka "
        "(IRCSL), comprises over 25 licensed insurance companies offering a broad spectrum of life, health, "
        "and general insurance products. Despite this breadth of supply, insurance penetration in Sri Lanka "
        "remains below 2% of GDP\u2014significantly lower than regional peers such as India (3.7%) and Malaysia "
        "(4.8%) (Swiss Re Institute, 2023). A primary barrier to higher penetration is the reliance on agent-led "
        "distribution, where customers must schedule face-to-face meetings with insurance advisors who manually "
        "assess needs, explain product features, and complete paper-based proposal forms."
    )
    rb.p(
        "This traditional process presents several inefficiencies: customers receive recommendations that may "
        "be influenced by agent commission structures rather than objective suitability; the absence of a "
        "standardised assessment framework means that two advisors may recommend different products for "
        "identical customer profiles; customers lack visibility into how recommendations are derived, "
        "undermining trust; and the manual, document-intensive proposal process creates friction that "
        "deters younger, digitally native demographics."
    )
    rb.p(
        "Insurance University is conceived as a digital-first alternative that addresses these pain points. "
        "The platform enables a customer to complete the entire insurance discovery and application process "
        "online\u2014from initial data capture through to receiving AI-generated recommendations, comparing "
        "products side by side, uploading KYC documents, and submitting a proposal\u2014without requiring any "
        "human intermediary. The system\u2019s scoring logic is transparent, explainable, and free from commission "
        "bias, thereby fostering customer trust and enabling informed decision-making."
    )
    rb.p(
        "The administrative stakeholder (insurance company product manager or underwriting team) uses the "
        "platform\u2019s admin console to manage the product catalogue, define underwriting rules, upload "
        "training data to refine the AI model, and monitor customer interaction patterns. This dual-purpose "
        "design ensures that the platform delivers value to both the demand side (customers seeking suitable "
        "insurance) and the supply side (insurers seeking efficient, data-driven distribution)."
    )

    rb.h2("2b  Goals of the Project")
    rb.p("Aim", bold=True)
    rb.p(
        "To design, develop, and evaluate an AI-powered web-based insurance recommendation platform that "
        "automates personalised insurance product assessment, providing Sri Lankan consumers with transparent, "
        "explainable, and bias-free product recommendations through a guided digital journey."
    )
    rb.p("Objectives", bold=True)
    rb.numbered(
        "Design and implement a multi-step wizard interface that captures customer personal, financial, "
        "and lifestyle data in a structured, user-friendly format."
    )
    rb.numbered(
        "Develop an AI scoring engine incorporating CART-based underwriting rules and collaborative "
        "filtering heuristics to rank insurance products by suitability, affordability, and risk alignment."
    )
    rb.numbered(
        "Implement a trainable model architecture that enables administrators to upload historical policy "
        "outcome data and retrain the scoring engine to improve recommendation accuracy over time."
    )
    rb.numbered(
        "Build a role-based administration console for managing insurance products, eligibility rules, "
        "training datasets, AI model versions, event logs, and customer insight analytics."
    )
    rb.numbered(
        "Provide real-time premium estimation with transparent breakdowns, affordability scoring, and "
        "lapse probability prediction to support informed customer decision-making."
    )
    rb.numbered(
        "Implement a secure KYC document upload workflow with version tracking and document reuse "
        "capabilities to streamline the proposal submission process."
    )
    rb.numbered(
        "Deliver the system as a containerised three-tier microservices architecture (Angular, Spring Boot, "
        "FastAPI) deployable via Docker Compose for consistent development, testing, and production environments."
    )

    # ── SECTION 3 ─────────────────────────────────────────────────────────
    rb.h1("3  The Scope of the Work")

    rb.h2("3a  The Current Situation and Context")
    rb.p(
        "The current insurance advisory process in Sri Lanka operates predominantly through a network of "
        "licensed insurance agents and brokers. A typical customer journey involves the following steps: "
        "the customer contacts an agent (often through a referral or walk-in visit); the agent conducts a "
        "verbal needs assessment; the agent manually selects one or two products from a printed product "
        "brochure; the customer receives a paper proposal form; the customer gathers supporting documents "
        "(NIC copies, medical reports, income certificates) and submits them physically; and the insurer\u2019s "
        "underwriting team reviews the application over a period of days to weeks."
    )
    rb.p("The sub-problems identified in this traditional process are as follows:", bold=True)
    rb.p(
        "Problem 1 \u2013 Lack of Objective Assessment: Agent recommendations may be influenced by commission "
        "incentives rather than genuine product suitability. Customers have no independent mechanism to "
        "verify whether a recommended product aligns with their risk profile and financial capacity."
    )
    rb.p(
        "Problem 2 \u2013 Inconsistent Recommendations: Different agents may recommend entirely different products "
        "for identical customer profiles, as there is no standardised scoring framework applied consistently "
        "across advisors."
    )
    rb.p(
        "Problem 3 \u2013 Opaque Pricing and Suitability Logic: Customers are unable to understand how premiums "
        "are calculated or why a particular product was recommended. This opacity undermines trust and "
        "discourages uptake."
    )
    rb.p(
        "Problem 4 \u2013 Manual, Document-Heavy Process: The paper-based proposal submission requires multiple "
        "physical visits and document handling, creating friction and delays that deter potential customers."
    )
    rb.p(
        "Problem 5 \u2013 Limited Data-Driven Product Development: Insurers lack systematic visibility into "
        "customer needs that go unmet by their current product portfolio, missing opportunities for "
        "product innovation."
    )
    rb.p(
        "Problem 6 \u2013 No Audit Trail or Analytics: Traditional advisory interactions are unrecorded, "
        "making it impossible to analyse customer journeys, identify drop-off points, or measure "
        "recommendation effectiveness."
    )

    rb.insert_image(os.path.join(DIAGRAM_DIR, "02_fishbone.png"), "Figure 2: Fishbone (Ishikawa) Diagram \u2013 Insurance Advisory Problems")
    rb.p(
        "The fishbone diagram above illustrates the root causes of low insurance penetration in Sri Lanka, "
        "categorised across six dimensions: People (agent bias, limited training), Processes (manual assessment, "
        "paper proposals), Technology (no digital platform, no AI scoring), Information (opaque pricing, no "
        "analytics), Environment (regulatory complexity, low digital literacy), and Products (limited "
        "customisation, complex feature sets)."
    )

    rb.h2("3b  Competing Products \u2013 Related Work")
    rb.p(
        "Several existing platforms address aspects of the insurance recommendation problem, though none "
        "provides the complete end-to-end solution that Insurance University targets. The following table "
        "summarises key competitors and their capabilities:"
    )
    rb.p("PolicyBazaar (India)", bold=True)
    rb.p(
        "PolicyBazaar is India\u2019s largest insurance aggregator platform, offering comparison tools across "
        "life, health, motor, and travel insurance. It provides premium calculators and allows users to "
        "purchase policies online. However, its recommendation engine relies primarily on rule-based filters "
        "rather than AI-driven scoring, and it does not incorporate underwriting logic or affordability "
        "analysis into its product ranking. The platform is not localised for the Sri Lankan market."
    )
    rb.p("Coverfox (India)", bold=True)
    rb.p(
        "Coverfox offers a similar aggregation model with a focus on health and motor insurance. It provides "
        "a chatbot-assisted advisory flow but lacks explainable AI scoring, premium breakdown transparency, "
        "and KYC document management within the platform."
    )
    rb.p("AIA Sri Lanka Digital Portal", bold=True)
    rb.p(
        "AIA\u2019s digital portal allows existing policyholders to view policy details and make payments online. "
        "However, it does not offer a product recommendation engine, wizard-based data capture, or any "
        "AI-driven advisory capability. New customer acquisition remains agent-driven."
    )
    rb.p("Ceylinco Life Mobile App", bold=True)
    rb.p(
        "Ceylinco Life\u2019s mobile application provides policy servicing features for existing customers but "
        "does not support product discovery or recommendation. The app focuses on claims tracking, premium "
        "payment reminders, and policy document access."
    )
    rb.p(
        "Insurance University differentiates itself from these competitors through: (a) a multi-stage AI "
        "scoring pipeline that combines underwriting rules with collaborative filtering and trainable weights; "
        "(b) transparent, explainable recommendations with detailed premium breakdowns; (c) an integrated "
        "KYC document upload and version management workflow; (d) an administration console with AI model "
        "lifecycle management; and (e) localisation for the Sri Lankan insurance market with LKR-denominated "
        "calculations and IRCSL-aligned product structures."
    )

    # ── SECTION 4 ─────────────────────────────────────────────────────────
    rb.h1("4  Stakeholders")
    rb.p(
        "The stakeholder analysis for Insurance University is presented using the Onion Model (Alexander & "
        "Robertson, 2004), which classifies stakeholders into concentric layers based on their proximity "
        "to the system."
    )
    rb.insert_image(os.path.join(DIAGRAM_DIR, "03_onion_model.png"), "Figure 3: Onion Model \u2013 Stakeholder View")

    rb.p("Layer 1 \u2013 The Product (Innermost)", bold=True)
    rb.p(
        "The system itself: the Angular frontend, Spring Boot backend, FastAPI AI engine, and MySQL database. "
        "All components must function correctly and interact seamlessly."
    )
    rb.p("Layer 2 \u2013 Direct Users", bold=True)
    rb.bullet(
        "Customers (Insurance Seekers): End-users who complete the wizard, receive recommendations, "
        "compare products, and submit proposals. They interact directly with the frontend."
    )
    rb.bullet(
        "Administrators (Product Managers / Underwriters): Insurance company staff who manage products, "
        "rules, training data, and models through the admin console."
    )
    rb.p("Layer 3 \u2013 Indirect Users and Operators", bold=True)
    rb.bullet(
        "IT Operations Team: Responsible for deploying, monitoring, and maintaining the Docker Compose "
        "infrastructure, database backups, and log management."
    )
    rb.bullet(
        "Data Science / AI Team: Responsible for preparing training datasets, evaluating model performance, "
        "and deciding when to promote new models."
    )
    rb.p("Layer 4 \u2013 The Wider Environment", bold=True)
    rb.bullet(
        "Insurance Regulatory Commission of Sri Lanka (IRCSL): The regulatory body whose guidelines govern "
        "product definitions, underwriting standards, and consumer protection requirements."
    )
    rb.bullet(
        "Government of Sri Lanka: Policy-makers interested in increasing insurance penetration and financial "
        "inclusion through digital platforms."
    )
    rb.bullet(
        "Banking and Financial Institutions: Potential partners for premium payment integration and "
        "customer referral pipelines."
    )
    rb.bullet(
        "Device and Infrastructure Providers: Cloud hosting providers, browser vendors, and mobile device "
        "manufacturers whose platforms the system runs on."
    )

    # ── SECTION 5 ─────────────────────────────────────────────────────────
    rb.h1("5  Constraints")

    rb.h2("5a  Implementation Environment of the Current System")
    rb.p(
        "Insurance University is designed as a locally deployable, containerised application stack. The "
        "implementation environment comprises the following components:"
    )
    rb.p("Table 1: Technology Stack Summary", bold=True, italic=True)
    rb.table(
        ["Component", "Technology", "Version", "Port"],
        [
            ["Frontend SPA", "Angular + Angular Material + Tailwind CSS", "Angular 21", "4200"],
            ["Backend API", "Spring Boot (Java, Maven)", "3.3.5 / Java 17", "8080"],
            ["AI Engine", "FastAPI + Uvicorn + Pydantic", "Python 3.11+", "8000"],
            ["Agentic Chat LLM", "Groq API (meta-llama/llama-4-scout-17b)", "OpenAI SDK", "\u2014"],
            ["Vector Memory", "FAISS + sentence-transformers", "all-MiniLM-L6-v2", "\u2014"],
            ["Database", "MySQL", "8.4 (16 tables)", "3306"],
            ["Orchestration", "Docker Compose", "Latest", "\u2014"],
            ["Authentication", "JWT (JJWT library, HMAC-SHA256)", "\u2014", "\u2014"],
            ["Streaming", "Server-Sent Events (SSE)", "\u2014", "\u2014"],
        ],
    )
    rb.p("")
    rb.insert_image(os.path.join(DIAGRAM_DIR, "04_high_level_arch.png"), "Figure 4: High-Level Architecture Diagram")
    rb.p(
        "The architecture follows a four-tier pattern. The Angular SPA communicates with the Spring Boot "
        "API via REST over HTTP, attaching JWT Bearer tokens for authenticated routes. For the agentic chat "
        "feature, Server-Sent Events (SSE) provide real-time streamed responses. The Spring Boot API "
        "communicates with the FastAPI AI engine via REST for scoring, training, and chat operations. The AI engine "
        "connects to the Groq LLM API for natural-language processing and uses FAISS for vector-based long-term "
        "memory storage. The MySQL "
        "database (16 tables) stores all persistent data including user accounts, customer sessions, answers, documents, "
        "products, rules, logs, chat messages, and model metadata. The AI engine persists trained model artefacts as JSON "
        "files and FAISS indices on the local filesystem."
    )
    rb.p(
        "For local development, Docker Compose orchestrates MySQL and the AI engine as containerised services, "
        "while the backend and frontend are run natively using Maven and Angular CLI respectively. The full "
        "stack can also be run entirely via Docker Compose for integration testing."
    )

    rb.h2("5b  Partner or Collaborative Applications")
    rb.p(
        "Insurance University integrates with one external AI inference service and is designed "
        "to accommodate additional partner integrations in the future."
    )
    rb.p("Active Integration", bold=True)
    rb.bullet(
        "Groq Cloud API (meta-llama/llama-4-scout-17b-16e-instruct): The agentic open chat subsystem "
        "uses the Groq inference platform via the OpenAI-compatible SDK for LLM-powered conversational "
        "AI. Groq provides high-speed inference for the ReAct tool-calling loop, enabling real-time "
        "streaming responses. The integration supports tool calling, structured outputs, and token tracking."
    )
    rb.bullet(
        "DuckDuckGo Search API: The agentic chat\u2019s web_search tool queries DuckDuckGo for real-time "
        "information about insurance regulations, market trends, and external topics not covered by the "
        "internal knowledge base."
    )
    rb.p("Planned Future Integrations", bold=True)
    rb.bullet(
        "Payment Gateway (e.g., PayHere, LankaPay): For online premium payment collection during the "
        "proposal submission stage."
    )
    rb.bullet(
        "IRCSL Regulatory API: For automated product registration validation and compliance reporting."
    )
    rb.bullet(
        "Hospital Information Systems: For automated retrieval of medical records to streamline the "
        "KYC verification process."
    )
    rb.bullet(
        "National Identity Verification Service: For automated NIC validation and identity confirmation "
        "using OCR or government APIs."
    )
    rb.bullet(
        "Email/SMS Notification Service: For customer notifications regarding recommendation availability, "
        "document upload reminders, and policy status updates."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # PART II: REQUIREMENTS
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("II  Requirements")

    # ── SECTION 6 ─────────────────────────────────────────────────────────
    rb.h1("6  Product Use Cases")

    rb.h2("6a  Use Case Diagrams")
    rb.p(
        "The system supports three primary actors: Customer, Administrator, and AI Engine. The Customer actor interacts "
        "with the public-facing journey (registration, wizard completion, recommendation viewing, product "
        "comparison, document upload, proposal submission, and agentic chat). The Administrator actor interacts with the "
        "back-office console (product management, rule management, training dataset upload, model lifecycle "
        "management, log analysis, and unmatched needs tracking). The AI Engine actor provides scoring and "
        "agentic chat capabilities powered by a Groq-hosted LLM with five specialised tools."
    )
    rb.insert_image(os.path.join(DIAGRAM_DIR, "05_use_case.png"), "Figure 5: Use Case Diagram")
    rb.p("The use case diagram illustrates the following use cases grouped by actor:")
    rb.p("Customer Use Cases:", bold=True)
    rb.bullet("UC-01: Register Account")
    rb.bullet("UC-02: Login to Platform")
    rb.bullet("UC-03: Complete Wizard (Step 1 \u2013 Personal, Step 2 \u2013 Financial, Step 3 \u2013 Preferences)")
    rb.bullet("UC-04: Get Insurance Recommendations")
    rb.bullet("UC-05: Compare Recommended Products")
    rb.bullet("UC-06: Use Premium Simulator")
    rb.bullet("UC-07: Upload KYC Documents")
    rb.bullet("UC-08: Review and Submit Proposal")
    rb.bullet("UC-09: View Past Sessions and Recommendations")
    rb.bullet("UC-10: Submit Feedback Survey")
    rb.bullet("UC-11: Wizard Chat (Chat-Based Data Capture)")
    rb.bullet("UC-12: Open Chat \u2013 Agentic AI Insurance Advisory (with 5 tools: KB, DB, Web, Calculator, Scoring)")

    rb.p("Administrator Use Cases:", bold=True)
    rb.bullet("UC-13: Admin Login")
    rb.bullet("UC-14: Manage Insurance Products (CRUD)")
    rb.bullet("UC-15: Manage Insurance Categories")
    rb.bullet("UC-16: Manage Eligibility Rules (CRUD)")
    rb.bullet("UC-17: Upload Training Dataset")
    rb.bullet("UC-18: Manage AI Model Versions")
    rb.bullet("UC-19: View and Export Event Logs")
    rb.bullet("UC-20: Manage Unmatched Customer Needs")
    rb.bullet("UC-21: Manage Pricing Tables")

    rb.h2("6b  Individual Product Use Cases")
    rb.p("Use Case 1: Get Insurance Recommendations (UC-04)", bold=True)
    rb.table(
        ["Element", "Description"],
        [
            ["Use Case ID", "UC-04"],
            ["Use Case Name", "Get Insurance Recommendations"],
            ["Primary Actor", "Customer"],
            ["Preconditions",
             "1. Customer is registered and logged in.\n"
             "2. Customer has an active session.\n"
             "3. Customer has completed Wizard Steps 1, 2, and 3 (all answers saved)."],
            ["Main Success Scenario",
             "1. Customer completes Wizard Step 3 and clicks \u2018Get Recommendations\u2019.\n"
             "2. Frontend sends POST /api/customer/sessions/{sessionId}/recommendations.\n"
             "3. Backend retrieves all saved answers for the session.\n"
             "4. Backend deserialises answer values and builds a feature map.\n"
             "5. Backend retrieves all active products from the database.\n"
             "6. Backend applies eligibility rules (age gates, smoker constraints, effective dates) to filter products.\n"
             "7. Backend serialises filtered products and features into a scoring request.\n"
             "8. Backend calls AiEngineClient.score() \u2192 POST /score on AI engine.\n"
             "9. AI engine runs CART underwriting for each product (age, diseases, clinical history, PEP/criminal).\n"
             "10. AI engine calculates suitability score (0.0\u20131.0) using collaborative filtering heuristics.\n"
             "11. AI engine calculates monthly premium estimate, affordability score, and lapse probability.\n"
             "12. AI engine predicts recommended coverage amount based on protection purpose.\n"
             "13. AI engine ranks products by score descending and generates follow-up questions.\n"
             "14. Backend persists full request/response as a RecommendationRunEntity (audit trail).\n"
             "15. Backend writes RECOMMENDATIONS_FETCHED event to session log.\n"
             "16. Backend returns ranked products and follow-up questions to frontend.\n"
             "17. Frontend displays ranked product cards with scores, premiums, reasons, and eligibility status."],
            ["Alternative Flows",
             "A1. AI Engine Unavailable: Backend returns 503 with error message; frontend shows retry option.\n"
             "A2. No Eligible Products: All products filtered by eligibility rules; AI engine receives empty list; "
             "response returns empty rankedProducts array; frontend shows \u2018no matching plans\u2019 message.\n"
             "A3. Underwriting Rejection: CART rules return \u2018No Offer\u2019 for all products; frontend displays "
             "ineligibility reasons."],
            ["Postconditions",
             "1. RecommendationRunEntity persisted with full request/response JSON.\n"
             "2. SessionLogEntity created with event type RECOMMENDATIONS_FETCHED.\n"
             "3. Customer can view and compare recommended products."],
            ["Non-Functional Requirements",
             "Response time < 5 seconds for typical scoring (5 products).\n"
             "Scoring must be deterministic for the same input."],
        ],
    )
    rb.p("")
    rb.p("Use Case 2: Upload Training Dataset and Train Model (UC-16 + UC-17)", bold=True)
    rb.table(
        ["Element", "Description"],
        [
            ["Use Case ID", "UC-16 / UC-17"],
            ["Use Case Name", "Upload Training Dataset and Train AI Model"],
            ["Primary Actor", "Administrator"],
            ["Preconditions",
             "1. Administrator is logged in with ADMIN role.\n"
             "2. A valid CSV file in policy-recommendation-v2 format is available.\n"
             "3. AI engine service is running and healthy."],
            ["Main Success Scenario",
             "1. Administrator navigates to Admin \u2192 Training page.\n"
             "2. Administrator clicks \u2018Upload Dataset\u2019 and selects a CSV file.\n"
             "3. Frontend sends POST /api/admin/training/datasets with multipart form data.\n"
             "4. Backend validates CSV format: checks for required columns (product_code, category_code, "
             "subcategory_code, policy_type, eligibility, outcome_score, age, smoker, income, "
             "monthlyexpenseslkr, networthlkr, conditions_count).\n"
             "5. Backend stores CSV file in uploads/ directory with UUID-prefixed filename.\n"
             "6. Backend creates DatasetMetaEntity with file metadata.\n"
             "7. Backend calls AiEngineClient.train() \u2192 POST /train on AI engine with CSV file.\n"
             "8. AI engine parses CSV rows and groups by attributes (age buckets, smoker status, product, category).\n"
             "9. AI engine calculates average outcome_score per group and derives weight adjustments.\n"
             "10. AI engine generates policy adjustment deltas per product/category/subcategory.\n"
             "11. AI engine persists model artefact as JSON file in data/models/ directory.\n"
             "12. AI engine returns training summary (rowsProcessed, updatedWeights, modelArtifactId).\n"
             "13. Backend creates ModelVersionEntity linked to the dataset.\n"
             "14. Backend returns combined response to frontend.\n"
             "15. Administrator views training results and clicks \u2018Promote\u2019 to activate the new model.\n"
             "16. Backend calls POST /models/{artifactId}/activate on AI engine.\n"
             "17. AI engine loads the specified model artefact as the active scoring model."],
            ["Alternative Flows",
             "A1. Invalid CSV Format: Backend or AI engine rejects file; error message returned.\n"
             "A2. Training Failure: AI engine returns error; Backend returns 500; admin sees error notification.\n"
             "A3. Retrain from Existing Dataset: Admin clicks \u2018Retrain\u2019 on a previously uploaded dataset; "
             "backend re-submits stored file to AI engine."],
            ["Postconditions",
             "1. DatasetMetaEntity persisted with file metadata.\n"
             "2. ModelVersionEntity created and optionally promoted to active.\n"
             "3. AI engine uses updated weights for subsequent scoring calls."],
            ["Non-Functional Requirements",
             "Training should complete within 30 seconds for datasets up to 10,000 rows.\n"
             "Model artefacts must be persisted for auditability."],
        ],
    )

    # ── SECTION 7 ─────────────────────────────────────────────────────────
    rb.add_page_break()
    rb.h1("7  Functional Requirements")
    rb.p(
        "The following table presents the functional requirements of Insurance University, cross-referenced "
        "with the use cases they support."
    )
    rb.p("Table 2: Functional Requirements", bold=True, italic=True)
    rb.table(
        ["ID", "Requirement", "Priority", "Related UC"],
        [
            ["FR-01", "The system shall allow customers to register with name, email, and password.", "Must", "UC-01"],
            ["FR-02", "The system shall authenticate customers using JWT tokens with 24-hour expiry.", "Must", "UC-02"],
            ["FR-03", "The system shall provide a multi-step wizard (3 steps) for capturing customer data.", "Must", "UC-03"],
            ["FR-04", "Step 1 shall capture personal data: age, smoker status, medical conditions, dependents.", "Must", "UC-03"],
            ["FR-05", "Step 2 shall capture financial data: monthly income, expenses, net worth, occupation hazard.", "Must", "UC-03"],
            ["FR-06", "Step 3 shall capture preferences: protection purpose, desired coverage, policy term, priority sliders.", "Must", "UC-03"],
            ["FR-07", "The system shall persist wizard answers as JSON key-value pairs per session.", "Must", "UC-03"],
            ["FR-08", "The system shall generate ranked product recommendations using an AI scoring engine.", "Must", "UC-04"],
            ["FR-09", "Recommendations shall include suitability score, monthly premium, affordability score, and reasons.", "Must", "UC-04"],
            ["FR-10", "The system shall apply eligibility rules to filter products before AI scoring.", "Must", "UC-04"],
            ["FR-11", "The system shall allow side-by-side comparison of selected recommended products.", "Should", "UC-05"],
            ["FR-12", "The system shall provide a premium simulator for what-if analysis.", "Should", "UC-06"],
            ["FR-13", "The system shall allow customers to upload KYC documents (NIC, medical, income proof).", "Must", "UC-07"],
            ["FR-14", "Document uploads shall support versioning with latest-version pointers.", "Should", "UC-07"],
            ["FR-15", "The system shall allow customers to review and complete proposal submissions.", "Must", "UC-08"],
            ["FR-16", "The system shall allow customers to view past sessions and recommendation history.", "Should", "UC-09"],
            ["FR-17", "The system shall provide CRUD operations for insurance products via admin console.", "Must", "UC-13"],
            ["FR-18", "The system shall provide CRUD operations for eligibility rules via admin console.", "Must", "UC-15"],
            ["FR-19", "The system shall allow administrators to upload CSV training datasets.", "Must", "UC-16"],
            ["FR-20", "The system shall automatically trigger AI model training upon dataset upload.", "Must", "UC-16"],
            ["FR-21", "The system shall allow administrators to promote trained models to active status.", "Must", "UC-17"],
            ["FR-22", "The system shall provide searchable, filterable, and exportable event logs.", "Should", "UC-18"],
            ["FR-23", "The system shall track unmatched customer needs for product gap analysis.", "Should", "UC-19"],
            ["FR-24", "The system shall persist full AI request/response payloads for audit trail.", "Must", "UC-04"],
            ["FR-25", "The system shall hash session identifiers (SHA-256) before writing to event logs.", "Must", "UC-19"],
            ["FR-26", "The system shall provide an agentic open chat interface with SSE-streamed real-time responses.", "Must", "UC-12"],
            ["FR-27", "The agentic chat shall support five tools: knowledge base, database query, web search, calculator, and policy scoring.", "Must", "UC-12"],
            ["FR-28", "The system shall maintain long-term conversational memory using FAISS vector storage.", "Should", "UC-12"],
            ["FR-29", "The agentic chat shall persist all messages (user, assistant, tool) in the open_chat_messages table.", "Must", "UC-12"],
            ["FR-30", "The wizard chat shall support both deterministic and LLM-based extraction modes for data capture.", "Should", "UC-11"],
            ["FR-31", "The system shall enforce configurable per-user daily and monthly token limits for open chat.", "Must", "UC-12"],
        ],
    )

    # ── SECTION 8 ─────────────────────────────────────────────────────────
    rb.add_page_break()
    rb.h1("8  Data Requirements")
    rb.p(
        "Insurance University manages a rich data model comprising 16 JPA entities that capture user identity, "
        "customer journey state, product catalogues, eligibility rules, AI model metadata, chat messages, and audit logs. "
        "The following table summarises the core entities and their primary data attributes."
    )
    rb.p("Table 4: JPA Entity Summary", bold=True, italic=True)
    rb.table(
        ["Entity", "Table Name", "Key Fields", "Relationships"],
        [
            ["UserEntity", "users", "id, name, email (unique), passwordHash, role (USER/ADMIN), createdAt", "\u2014"],
            ["CustomerSessionEntity", "customer_sessions", "id (UUID), status, userEmail, createdAt, updatedAt", "1:N Answers, Documents, Runs, Chat"],
            ["CustomerAnswerEntity", "customer_answers", "id, sessionId (FK), key, valueJson (JSON), createdAt", "N:1 Session"],
            ["CustomerDocumentEntity", "customer_documents", "id, sessionId, userEmail, docType, docSide, storedPath, versionNo, latestForSession", "N:1 Session"],
            ["CustomerChatMessageEntity", "customer_chat_messages", "id, sessionId, role (USER/AGENT), message (TEXT), metadataJson", "N:1 Session"],
            ["OpenChatMessage", "open_chat_messages", "id, sessionId, userId, role (USER/ASSISTANT/TOOL), content (TEXT), toolName, toolArgs (JSON), toolResult (TEXT), tokensUsed, createdAt", "\u2014"],
            ["RecommendationRunEntity", "recommendation_runs", "id, sessionId, requestJson, responseJson, createdAt", "N:1 Session"],
            ["SessionLogEntity", "session_logs", "id, sessionHash (SHA-256), timestamp, eventType, userSegment, payloadJson", "\u2014"],
            ["InsuranceCategoryEntity", "insurance_categories", "id, code, name, description, active, displayOrder", "1:N Subcategories, Products"],
            ["InsuranceSubcategoryEntity", "insurance_subcategories", "id, code, name, description, active, category_id (FK)", "N:1 Category, 1:N Products"],
            ["ProductEntity", "products", "id, code (unique), name, basePremium, tagsJson, eligibilityJson, ridersJson, benefitsJson", "N:1 Category, N:1 Subcategory"],
            ["EligibilityRuleEntity", "eligibility_rules", "id, name, ruleJson, version, effectiveFrom, effectiveTo", "\u2014"],
            ["DatasetMetaEntity", "dataset_meta", "id, originalFilename, storedPath, fileSize, formatVersion, trainingGoal", "1:N Models"],
            ["ModelVersionEntity", "model_versions", "id, name, description, artifactId, sourceDatasetId (FK), rowsProcessed, active", "N:1 Dataset"],
            ["PricingTableEntity", "pricing_tables", "id, name, pricingJson, version, effectiveFrom, effectiveTo", "\u2014"],
            ["UnmatchedNeedEntity", "unmatched_needs", "id, theme, occurrences, sampleAnonymisedText, updatedAt", "\u2014"],
        ],
    )

    rb.h2("8a  Safety, Ethical, and Critical Requirements")
    rb.p(
        "Insurance recommendation systems handle sensitive personal and financial data and produce outputs "
        "that directly influence financial decisions. The following safety and ethical requirements are "
        "therefore critical:"
    )
    rb.numbered(
        "Non-Discriminatory Scoring: The AI scoring engine must not discriminate based on ethnicity, "
        "religion, gender, or other protected characteristics. The scoring pipeline uses only "
        "actuarially relevant factors (age, smoker status, health conditions, financial capacity)."
    )
    rb.numbered(
        "Explainable Recommendations: Every product recommendation must include human-readable reasons "
        "explaining why the product was scored highly or poorly. This transparency is essential for "
        "customer trust and regulatory compliance."
    )
    rb.numbered(
        "Data Minimisation: The system collects only the data necessary for insurance assessment. "
        "No unnecessary personal identifiers are stored in event logs (session IDs are hashed with SHA-256)."
    )
    rb.numbered(
        "Consent-Based Data Handling: Customers can delete their sessions at any time, which removes "
        "all associated answers, documents, and recommendation runs. This supports the right to erasure."
    )
    rb.numbered(
        "Audit Trail Integrity: All AI scoring calls are persisted with full request and response payloads "
        "in the RecommendationRunEntity table, enabling post-hoc review of any recommendation made."
    )
    rb.numbered(
        "No Automated Decision-Making: The system provides recommendations, not binding decisions. "
        "Final product selection and policy issuance remain human-controlled processes."
    )

    # ── SECTION 9 ─────────────────────────────────────────────────────────
    rb.add_page_break()
    rb.h1("9  Security Requirements")

    rb.h2("9a  Access Requirements")
    rb.p(
        "Insurance University implements a role-based access control (RBAC) model with two roles: "
        "CUSTOMER (USER) and ADMIN. Access is enforced at the API layer through Spring Security\u2019s "
        "filter chain and at the frontend through Angular route guards."
    )
    rb.bullet(
        "Public Endpoints: Authentication routes (/api/auth/register, /api/auth/login, /api/auth/admin/login) "
        "and client-side event logging (/api/logs) are accessible without authentication."
    )
    rb.bullet(
        "Customer Endpoints: All /api/customer/** routes require a valid JWT token with USER role. "
        "Session-level access control ensures that customers can only access their own sessions, "
        "answers, documents, and recommendations."
    )
    rb.bullet(
        "Admin Endpoints: All /api/admin/** routes require a valid JWT token with ADMIN role. "
        "Admin access is restricted to a single pre-configured administrator account."
    )
    rb.bullet(
        "Frontend Guards: Angular route guards (customerAuthGuard, adminGuard, wizardProgressGuard) "
        "enforce client-side access control, redirecting unauthenticated users to appropriate login pages."
    )

    rb.h2("9b  Integrity Requirements")
    rb.p(
        "Data integrity is maintained through the following mechanisms:"
    )
    rb.bullet(
        "Database Integrity: MySQL enforces foreign key constraints between entities (e.g., session \u2192 "
        "answers, session \u2192 documents). JPA\u2019s ddl-auto: update mode manages schema evolution."
    )
    rb.bullet(
        "Session Hashing: Customer session identifiers are hashed using SHA-256 before being written to "
        "event logs, ensuring that raw session IDs are not exposed in analytics data."
    )
    rb.bullet(
        "Document Version Control: Each document upload creates a new version entry with incrementing "
        "version numbers. Latest-version pointers (latestForSession, latestForUser) are updated "
        "transactionally to prevent inconsistent state."
    )
    rb.bullet(
        "Recommendation Audit Trail: Full AI request and response payloads are persisted as JSON in "
        "the RecommendationRunEntity table, providing an immutable record of every recommendation generated."
    )
    rb.bullet(
        "Input Validation: All API inputs are validated at the controller level using Spring\u2019s "
        "@Valid annotations and custom validation logic."
    )

    rb.h2("9c  Privacy Requirements")
    rb.p(
        "Privacy is a core design consideration given the sensitive nature of insurance-related personal "
        "and financial data:"
    )
    rb.bullet(
        "Password Hashing: Customer passwords are hashed using BCrypt before storage. Raw passwords "
        "are never persisted or logged."
    )
    rb.bullet(
        "Log Redaction: The RequestResponseLoggingFilter automatically redacts sensitive fields "
        "(passwords, tokens, authorisation headers) from HTTP request/response logs."
    )
    rb.bullet(
        "Session Anonymisation: Event logs use SHA-256 hashed session identifiers rather than raw "
        "session UUIDs, preventing direct linkage between analytics data and individual customer sessions."
    )
    rb.bullet(
        "Document Access Control: KYC documents can only be downloaded by the customer who uploaded "
        "them (verified via JWT email claim against the session\u2019s userEmail)."
    )
    rb.bullet(
        "Data Deletion: Customers can delete sessions, which cascades to remove all associated answers, "
        "documents (both database records and physical files), and recommendation runs."
    )

    rb.h2("9d  Accessibility Requirements")
    rb.p(
        "The frontend is built with Angular Material, which provides built-in accessibility features "
        "aligned with WCAG 2.1 Level AA standards:"
    )
    rb.bullet(
        "Semantic HTML: Angular Material components render semantic HTML elements with appropriate "
        "ARIA attributes for screen reader compatibility."
    )
    rb.bullet(
        "Keyboard Navigation: All interactive elements (forms, buttons, navigation) are accessible "
        "via keyboard-only navigation with visible focus indicators."
    )
    rb.bullet(
        "Colour Contrast: Tailwind CSS utilities are used with contrast-compliant colour palettes "
        "that meet WCAG AA contrast ratio requirements (4.5:1 for normal text)."
    )
    rb.bullet(
        "Responsive Design: The application is responsive and functions correctly across desktop, "
        "tablet, and mobile viewport sizes."
    )
    rb.bullet(
        "Form Validation Feedback: Error messages are associated with form fields using ARIA "
        "attributes and are announced by screen readers."
    )

    # ── SECTION 10 ────────────────────────────────────────────────────────
    rb.h1("10  Legal Requirements")

    rb.h2("10a  Compliance Requirements")
    rb.p(
        "Insurance University must comply with the following legal and regulatory frameworks:"
    )
    rb.numbered(
        "Personal Data Protection Act (PDPA) of Sri Lanka: The system handles personal data including "
        "names, email addresses, financial information, health conditions, and identity documents. "
        "Compliance requires: lawful basis for processing (consent through registration), data minimisation, "
        "purpose limitation, storage limitation, and the right to erasure (implemented via session deletion)."
    )
    rb.numbered(
        "Regulation of Insurance Industry Act No. 43 of 2000 (IRCSL): Insurance products offered on "
        "the platform must be registered with IRCSL. Product definitions, eligibility criteria, and "
        "premium structures must align with the regulator\u2019s published guidelines. The admin console "
        "allows authorised administrators to maintain compliant product configurations."
    )
    rb.numbered(
        "Computer Crimes Act No. 24 of 2007 (Sri Lanka): The system must implement adequate security "
        "measures to prevent unauthorised access to customer data. This is addressed through JWT "
        "authentication, RBAC, BCrypt password hashing, and HTTPS (recommended for production deployment)."
    )
    rb.numbered(
        "Consumer Affairs Authority Act No. 9 of 2003: Product recommendations must be fair, "
        "transparent, and not misleading. The explainable AI approach, with visible scoring reasons "
        "and premium breakdowns, supports compliance with consumer protection standards."
    )
    rb.numbered(
        "Electronic Transactions Act No. 19 of 2006: Digital signatures and electronic records "
        "generated by the platform (e.g., proposal summaries, recommendation audit trails) are "
        "recognised as valid under Sri Lankan law, supporting the shift from paper-based to "
        "digital insurance processes."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # PART III: DESIGN
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("III  Design")

    # ── SECTION 11 ────────────────────────────────────────────────────────
    rb.h1("11  Proposed Software Architecture")

    rb.h2("11a  Class Diagrams")
    rb.p(
        "The system\u2019s persistent data model comprises 16 JPA entity classes organised across the "
        "following domain areas: User Authentication, Customer Journey, Product Catalogue, Rules Engine, "
        "Training/Model Management, Agentic Chat, and Analytics. The complete class diagram with attributes and "
        "operations is presented below."
    )
    rb.insert_image(os.path.join(DIAGRAM_DIR, "06_class_diagram.png"), "Figure 6: Class Diagram \u2013 Backend Entity Model (16 Entities)")
    rb.p("The key classes and their relationships are as follows:", bold=True)
    rb.p(
        "UserEntity represents the core identity for both customer and administrator users. It contains "
        "fields for name, email (unique constraint), passwordHash (BCrypt-encoded), role (enumerated as "
        "USER or ADMIN), and createdAt timestamp. Authentication operations (register, login) interact "
        "with this entity through the AuthController and AuthService."
    )
    rb.p(
        "CustomerSessionEntity is the central orchestrating entity for the customer journey. Each session "
        "has a UUID identifier, a status field (ACTIVE or COMPLETED), and a reference to the customer\u2019s "
        "email. It maintains one-to-many relationships with CustomerAnswerEntity (wizard responses), "
        "CustomerDocumentEntity (KYC uploads), RecommendationRunEntity (AI scoring results), and "
        "CustomerChatMessageEntity (conversational data capture messages)."
    )
    rb.p(
        "OpenChatMessage is a new entity introduced for the agentic chat subsystem. It stores all messages "
        "exchanged during open chat sessions, including user messages, assistant responses, and tool call "
        "results. Key fields include sessionId, userId, role (USER/ASSISTANT/TOOL), content (TEXT), "
        "toolName, toolArgs (JSON), toolResult (TEXT), tokensUsed (for rate limiting), and createdAt. "
        "This entity enables full auditability of AI-generated insurance advice and tool usage."
    )
    rb.p(
        "ProductEntity represents an insurance product within a hierarchical catalogue. It belongs to "
        "an InsuranceCategoryEntity (e.g., Life Insurance, Retirement Plans) and an "
        "InsuranceSubcategoryEntity (e.g., Protection Plans). Product attributes include code, name, "
        "basePremium, tagsJson (array), benefitsJson, ridersJson, eligibilityJson, and age/term "
        "constraints (minEligibleAge, maxEligibleAge, minPolicyTermYears, maxPolicyTermYears)."
    )
    rb.p(
        "EligibilityRuleEntity stores administrator-defined business rules as JSON objects containing "
        "minAge, maxAge, and smokerAllowed constraints. Rules are versioned and have effective date "
        "ranges, enabling temporal rule management."
    )
    rb.p(
        "DatasetMetaEntity and ModelVersionEntity together manage the AI model lifecycle. DatasetMetaEntity "
        "stores metadata about uploaded training CSV files, while ModelVersionEntity tracks each trained "
        "model\u2019s name, artefact identifier, source dataset, row count, and active status. Only one "
        "model may be active at any time."
    )
    rb.p(
        "SessionLogEntity provides the immutable event log. Each entry contains a SHA-256 hashed session "
        "identifier, timestamp, event type (enumerated), user segment, and a JSON payload. This entity "
        "supports the admin log viewer with full-text search and export capabilities."
    )

    rb.h2("11b  Dynamic Model")
    rb.p(
        "The following sequence diagram depicts the end-to-end flow for the \u2018Get Insurance "
        "Recommendations\u2019 use case (UC-04), the system\u2019s most complex interaction."
    )
    rb.insert_image(os.path.join(DIAGRAM_DIR, "07_sequence.png"), "Figure 7: Sequence Diagram \u2013 Get Insurance Recommendations")
    rb.p("Lifelines:", bold=True)
    rb.bullet("Customer (Actor)")
    rb.bullet("Angular Frontend (SPA)")
    rb.bullet("CustomerAuthInterceptor")
    rb.bullet("Spring Boot API (CustomerController)")
    rb.bullet("CustomerSessionService")
    rb.bullet("CustomerAnswerRepository")
    rb.bullet("ProductRepository")
    rb.bullet("EligibilityRuleRepository")
    rb.bullet("AiEngineClient")
    rb.bullet("FastAPI AI Engine")
    rb.bullet("RecommendationRunRepository")
    rb.bullet("SessionLogRepository")
    rb.p("Sequence of Messages:", bold=True)
    rb.numbered(
        "Customer clicks \u2018Get Recommendations\u2019 on Step 3 completion screen."
    )
    rb.numbered(
        "Angular SPA sends HTTP POST /api/customer/sessions/{sessionId}/recommendations."
    )
    rb.numbered(
        "CustomerAuthInterceptor attaches JWT Bearer token to the request header."
    )
    rb.numbered(
        "CustomerController receives request and delegates to CustomerSessionService.getRecommendations(sessionId)."
    )
    rb.numbered(
        "CustomerSessionService calls CustomerAnswerRepository.findBySessionId(sessionId) to retrieve all answers."
    )
    rb.numbered(
        "Service deserialises JSON valueJson fields and constructs a feature map (age, smoker, income, expenses, etc.)."
    )
    rb.numbered(
        "Service calls ProductRepository.findAll() to retrieve all active products."
    )
    rb.numbered(
        "Service calls EligibilityRuleRepository.findAll() and filters products by rule constraints (age, smoker, effective dates)."
    )
    rb.numbered(
        "Service serialises filtered products into a list of product maps with metadata (code, name, basePremium, tags, category, riders, etc.)."
    )
    rb.numbered(
        "Service calls AiEngineClient.score(sessionId, featureMap, productList)."
    )
    rb.numbered(
        "AiEngineClient sends HTTP POST /score to FastAPI AI Engine with JSON body."
    )
    rb.numbered(
        "AI Engine executes the scoring pipeline:\n"
        "  a. CART underwriting check per product (age gates, disease rules, clinical history, legal checks).\n"
        "  b. Collaborative filtering score calculation (0.0\u20131.0) per product.\n"
        "  c. Policy type classification per product.\n"
        "  d. Coverage prediction based on protection purpose.\n"
        "  e. Monthly premium calculation with risk multipliers and riders.\n"
        "  f. Affordability scoring and lapse probability estimation.\n"
        "  g. Rank products by score descending.\n"
        "  h. Generate follow-up questions based on data completeness and confidence."
    )
    rb.numbered(
        "AI Engine returns JSON response with rankedProducts array and followUpQuestions array."
    )
    rb.numbered(
        "AiEngineClient returns the response map to CustomerSessionService."
    )
    rb.numbered(
        "Service persists RecommendationRunEntity with requestJson and responseJson."
    )
    rb.numbered(
        "Service writes SessionLogEntity with event type RECOMMENDATIONS_FETCHED and hashed sessionId."
    )
    rb.numbered(
        "Service returns response map to CustomerController."
    )
    rb.numbered(
        "Controller returns HTTP 200 with JSON response body to Angular SPA."
    )
    rb.numbered(
        "Angular SPA renders recommendation cards showing score, premium, affordability, reasons, and eligibility for each product."
    )

    rb.h2("11c  System Model and Decomposition")
    rb.p(
        "The system model for Insurance University integrates the functional model (use cases), the object "
        "model (class diagram), and the dynamic model (sequence diagrams) into a consistent whole. The "
        "following compliance properties are demonstrated:"
    )
    rb.numbered(
        "Use-Case-Driven Class Identification: Every entity in the class diagram is traceable to one or "
        "more use cases. For example, CustomerSessionEntity supports UC-03 (Complete Wizard), UC-04 "
        "(Get Recommendations), and UC-07 (Upload Documents). ProductEntity supports UC-13 (Manage "
        "Products) and UC-04 (Get Recommendations). ModelVersionEntity supports UC-17 (Manage AI Models)."
    )
    rb.numbered(
        "Functional Requirement Coverage: The use case diagram captures the functional requirements from "
        "the end user\u2019s perspective. Each functional requirement (FR-01 through FR-25) maps to at "
        "least one use case, ensuring complete coverage."
    )
    rb.numbered(
        "Sequence-Use Case Alignment: The sequence diagram for UC-04 (Get Recommendations) illustrates "
        "the runtime flow of a single use case scenario, showing how objects interact at the instance "
        "level. This complements the class diagram\u2019s structural view."
    )
    rb.numbered(
        "Structural-Dynamic Consistency: Classes appearing in the class diagram (CustomerSessionService, "
        "AiEngineClient, ProductRepository, etc.) also appear as lifelines in the sequence diagram, "
        "demonstrating that the structural and behavioural models describe the same system."
    )

    rb.h2("11d  Persistent Data Management")
    rb.p(
        "The system manages two categories of data: persistent data that must be stored long-term for "
        "operational, audit, and regulatory purposes, and ephemeral data that exists only during "
        "transient processes."
    )
    rb.insert_image(os.path.join(DIAGRAM_DIR, "08_er_diagram.png"), "Figure 8: Entity-Relationship Diagram (16 Tables)")
    rb.p("Table 6: Persistent vs Ephemeral Data Classification", bold=True, italic=True)
    rb.table(
        ["Data Item", "Classification", "Storage Mechanism", "Rationale"],
        [
            ["User accounts (email, name, passwordHash, role)", "Persistent", "MySQL (users table)", "Required for authentication and identity management"],
            ["Customer session records", "Persistent", "MySQL (customer_sessions)", "Required for session history and audit trail"],
            ["Wizard answers (JSON key-value pairs)", "Persistent", "MySQL (customer_answers)", "Required for recommendation input and replay"],
            ["KYC documents (NIC, medical, income)", "Persistent", "Filesystem + MySQL metadata", "Required for proposal verification and regulatory compliance"],
            ["Recommendation runs (request/response JSON)", "Persistent", "MySQL (recommendation_runs)", "Required for audit trail and regulatory review"],
            ["Event logs (hashed session, event type, payload)", "Persistent", "MySQL (session_logs)", "Required for analytics, compliance, and debugging"],
            ["Open chat messages (user, assistant, tool)", "Persistent", "MySQL (open_chat_messages)", "Required for agentic chat audit trail and token usage tracking"],
            ["Insurance products, categories, subcategories", "Persistent", "MySQL (products, categories)", "Reference data for product catalogue"],
            ["Eligibility rules (JSON rule definitions)", "Persistent", "MySQL (eligibility_rules)", "Business rules for underwriting"],
            ["Training datasets (CSV files)", "Persistent", "Filesystem + MySQL metadata", "Required for model retraining and audit"],
            ["Trained model artefacts (JSON weight files)", "Persistent", "AI engine filesystem (data/models/)", "Required for scoring and model version management"],
            ["Model version records", "Persistent", "MySQL (model_versions)", "Required for model lifecycle governance"],
            ["Pricing tables", "Persistent", "MySQL (pricing_tables)", "Reference data for premium lookup"],
            ["Unmatched needs records", "Persistent", "MySQL (unmatched_needs)", "Product gap analysis data"],
            ["FAISS vector indices (chat memory)", "Persistent", "AI engine filesystem (data/memory/)", "Long-term conversational memory for agentic chat"],
            ["JWT tokens (in-transit)", "Ephemeral", "HTTP headers (client memory)", "Valid for 24 hours; not persisted server-side"],
            ["Wizard form state (in-progress)", "Ephemeral", "Angular service (browser memory)", "Held in WizardStateService until saved to backend"],
            ["AI engine runtime weights (in-memory)", "Ephemeral", "Python process memory", "Loaded from persisted model artefact at startup/activation"],
            ["Shopping cart / selected products for comparison", "Ephemeral", "Angular service (browser memory)", "Held during comparison workflow only"],
            ["HTTP request/response bodies (in-transit)", "Ephemeral", "Network layer", "Logged selectively by RequestResponseLoggingFilter"],
        ],
    )

    # ── SECTION 12 ────────────────────────────────────────────────────────
    rb.add_page_break()
    rb.h1("12  User Interface")
    rb.p(
        "This section presents the key user interface screens of Insurance University, organised by "
        "user journey. The frontend is built with Angular 21, Angular Material components, and Tailwind "
        "CSS for responsive styling."
    )
    rb.h2("Customer Journey Interfaces")

    rb.h3("Landing Page")
    rb.p(
        "The landing page provides an overview of the Insurance University platform with clear "
        "call-to-action buttons for starting the wizard or taking a survey. The design emphasises "
        "simplicity and trust-building with concise messaging about the AI-powered recommendation process."
    )
    rb.placeholder_image("Figure 9: Landing Page")

    rb.h3("Wizard Step 1 \u2013 Personal Information")
    rb.p(
        "The first wizard step captures the customer\u2019s personal and health data. Form fields include "
        "age (number input), smoker status (toggle), medical conditions (multi-select checkboxes for "
        "heart disease, cancer, stroke, kidney disorder, etc.), number of dependents (number input), "
        "and children\u2019s ages (comma-separated or individual inputs). Form validation ensures all "
        "required fields are completed before proceeding."
    )
    rb.placeholder_image("Figure 9: Wizard Step 1 \u2013 Personal Information")

    rb.h3("Wizard Step 2 \u2013 Financial Information")
    rb.p(
        "The second wizard step captures financial data including monthly income (LKR), monthly expenses "
        "(LKR), net worth (LKR), and occupation hazard level (1\u20135 slider). Angular Material form "
        "fields with currency formatting and validation are used. The wizard progress guard prevents "
        "access to Step 2 unless Step 1 is completed."
    )
    rb.placeholder_image("Figure 10: Wizard Step 2 \u2013 Financial Information")

    rb.h3("Wizard Step 3 \u2013 Insurance Preferences")
    rb.p(
        "The third wizard step captures insurance-specific preferences: protection purpose (dropdown: "
        "Survivor Income, Education Funding, Retirement Supplement, Estate Liquidity), desired sum "
        "assured (LKR input), desired policy term (years slider), and priority sliders for Safety "
        "(1\u20135) and Equity (1\u20135). Upon completion, the system automatically triggers the "
        "recommendation engine."
    )
    rb.placeholder_image("Figure 11: Wizard Step 3 \u2013 Insurance Preferences")

    rb.h3("Recommendations View")
    rb.p(
        "The recommendations view displays AI-generated product cards ranked by suitability score. "
        "Each card shows: product name, suitability score (0\u2013100%), monthly premium estimate (LKR), "
        "affordability score, eligibility decision (Eligible / No Offer / Referral Required), "
        "coverage recommendation (LKR), lapse probability, and a list of human-readable reasons "
        "explaining the score. Customers can select products for side-by-side comparison."
    )
    rb.placeholder_image("Figure 12: Recommendations View")

    rb.h3("Product Comparison View")
    rb.p(
        "The comparison view presents two or three selected products in a side-by-side table format. "
        "Comparison dimensions include: benefits, riders (with selection checkboxes), premium breakdown "
        "(base premium, coverage factor, risk multiplier, rider costs, tax), eligibility constraints, "
        "and affordability metrics. The \u2018Proceed to Proposal\u2019 button advances the journey."
    )
    rb.placeholder_image("Figure 13: Product Comparison View")

    rb.h3("Document Upload View")
    rb.p(
        "The document upload interface provides separate upload areas for NIC (front and back), "
        "medical report, and income proof. Each upload area shows the current upload status (pending, "
        "uploaded, or version number). Drag-and-drop and file browser selection are both supported. "
        "Previously uploaded documents (from other sessions) can be reused via the \u2018latest user "
        "documents\u2019 feature."
    )
    rb.placeholder_image("Figure 14: Proposal Upload View")

    rb.h2("Administrator Journey Interfaces")

    rb.h3("Admin Dashboard")
    rb.p(
        "The admin dashboard provides a summary of key performance indicators: total active sessions, "
        "recommendations generated, documents processed, and models trained. A recent activity timeline "
        "shows the latest events across the platform."
    )
    rb.placeholder_image("Figure 15: Admin Dashboard")

    rb.h3("Training Management")
    rb.p(
        "The training management page is divided into two sections: Datasets (upload CSV, view upload "
        "history, retrain from existing dataset) and Models (view trained models, active model indicator, "
        "promote model to active). Upload metadata includes file name, size, format version, upload date, "
        "and training results (rows processed, updated weights, skipped rows)."
    )
    rb.placeholder_image("Figure 16: Admin Training Management")

    rb.h3("Log Viewer")
    rb.p(
        "The log viewer provides a searchable, filterable table of system events. Filters include date "
        "range (from/to date pickers), event type (dropdown: SESSION_CREATED, ANSWERS_SUBMITTED, "
        "RECOMMENDATIONS_FETCHED, DOCUMENT_UPLOADED, SESSION_DELETED, SESSION_COMPLETED), session hash "
        "(text input), and user segment (text input). Export buttons allow downloading filtered results "
        "as CSV or JSON."
    )
    rb.placeholder_image("Figure 17: Admin Logs Viewer")

    rb.h3("Open Chat Interface (Agentic AI)")
    rb.p(
        "The open chat interface provides a conversational AI experience for insurance advisory. "
        "Customers can ask questions about insurance products, request personalised recommendations, "
        "and explore coverage options through natural language. The interface features:\n"
        "\u2022 Real-time SSE-streamed responses rendered token-by-token.\n"
        "\u2022 Tool call visualisation showing which tools the AI used (knowledge base, calculator, etc.).\n"
        "\u2022 Markdown-formatted responses for structured information display.\n"
        "\u2022 Conversation history persistence and session management.\n"
        "\u2022 Five specialised tools: insurance knowledge base, database query, web search, calculator, "
        "and policy scoring."
    )
    rb.placeholder_image("Figure 18: Open Chat (Agentic AI) Interface")

    # ═══════════════════════════════════════════════════════════════════════
    # PART IV: IMPLEMENTATION
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("IV  Implementation")

    # ── SECTION 13 ────────────────────────────────────────────────────────
    rb.h1("13  Coding")

    rb.h2("Tools and Technologies")
    rb.p(
        "Insurance University was developed using a modern, full-stack technology portfolio selected "
        "for productivity, performance, and maintainability. The following sections detail the key "
        "tools and their roles in the project."
    )
    rb.p("Frontend Development Environment", bold=True)
    rb.bullet("Angular 21: Component-based SPA framework with standalone components and lazy-loaded routes.")
    rb.bullet("Angular Material: Pre-built, accessible UI components (forms, tables, dialogs, navigation).")
    rb.bullet("Tailwind CSS: Utility-first CSS framework for responsive layout and custom styling.")
    rb.bullet("TypeScript: Statically typed superset of JavaScript for type safety.")
    rb.bullet("Angular CLI: Build, serve, test, and lint tooling.")
    rb.bullet("RxJS: Reactive programming library for asynchronous data streams and HTTP responses.")

    rb.p("Backend Development Environment", bold=True)
    rb.bullet("Java 17: Primary backend language with modern language features (records, pattern matching).")
    rb.bullet("Spring Boot 3.3.5: Convention-over-configuration framework for REST API development.")
    rb.bullet("Spring Security: Authentication and authorisation framework with JWT filter chain.")
    rb.bullet("Spring Data JPA: Object-relational mapping with Hibernate for MySQL persistence.")
    rb.bullet("JJWT (io.jsonwebtoken): JSON Web Token generation and validation library.")
    rb.bullet("BCrypt: Password hashing algorithm via Spring Security\u2019s PasswordEncoder.")
    rb.bullet("Maven: Dependency management and build automation.")
    rb.bullet("Log4j2: Structured logging with configurable output formats.")

    rb.p("AI Engine Development Environment", bold=True)
    rb.bullet("Python 3.11+: Primary language for the AI scoring engine.")
    rb.bullet("FastAPI: High-performance async web framework for REST API development.")
    rb.bullet("Uvicorn: ASGI server for serving the FastAPI application.")
    rb.bullet("Pydantic: Data validation and serialisation using Python type hints.")

    rb.p("Infrastructure and DevOps", bold=True)
    rb.bullet("Docker: Containerisation for MySQL and AI engine services.")
    rb.bullet("Docker Compose: Multi-container orchestration for development and integration testing.")
    rb.bullet("MySQL 8.4: Relational database with InnoDB storage engine.")
    rb.bullet("Git: Version control system for source code management.")

    rb.h2("Key Implementation Details")

    rb.h3("1. JWT Authentication Implementation")
    rb.p(
        "The authentication system is built on stateless JWT tokens using the JJWT library. The "
        "implementation consists of two core classes: JwtService for token lifecycle management and "
        "JwtAuthFilter for request-level authentication."
    )
    rb.p(
        "JwtService provides four operations: generateToken(subject, role) creates a signed JWT with "
        "the user\u2019s email as the subject, role as a custom claim, and a 24-hour expiration; "
        "extractSubject(token) parses the token and returns the email; extractRole(token) returns the "
        "role claim (defaulting to USER if absent); and isValid(token) verifies the signature and "
        "expiration. Tokens are signed using HMAC-SHA256 with a 256-bit secret key configured in "
        "application.yml."
    )
    rb.p(
        "JwtAuthFilter extends OncePerRequestFilter and intercepts every HTTP request. It extracts "
        "the Bearer token from the Authorization header, validates it using JwtService, and constructs "
        "a UsernamePasswordAuthenticationToken with the user\u2019s email as principal and ROLE_{role} as "
        "the granted authority. This authentication object is placed in the SecurityContext, enabling "
        "Spring Security\u2019s @PreAuthorize annotations and hasRole() checks in the SecurityConfig "
        "filter chain."
    )
    rb.p("This implementation is novel code, designed and developed specifically for this project.", italic=True)

    rb.h3("2. AI Scoring Pipeline")
    rb.p(
        "The AI engine\u2019s scoring pipeline (implemented in main.py) executes seven sequential tasks "
        "for each product in the recommendation request. This multi-stage approach is the project\u2019s "
        "most significant technical contribution."
    )
    rb.p("Task 1 \u2013 CART Underwriting Engine", bold=True)
    rb.p(
        "A Classification and Regression Tree (CART) implementation evaluates underwriting eligibility. "
        "The decision tree applies age gates (reject if age < 18 or \u2265 65.5), disease disclosure rules "
        "(reject if any critical condition is disclosed), clinical history checks (exclude Critical Illness "
        "rider if hospitalised within 5 years), and legal checks (refer if PEP or criminal history). "
        "Outputs include an eligibility decision (Eligible, No Offer, or Referral Required) and a list "
        "of exclusion reasons."
    )
    rb.p("Task 2 \u2013 Collaborative Filtering Scoring", bold=True)
    rb.p(
        "Each product receives a suitability score between 0.0 and 1.0 calculated through a weighted "
        "combination of heuristic rules and learned adjustments. Base heuristics include: young age boost "
        "(+0.10 if age < 30), old age penalty (\u20130.15 if age > 55), smoker penalty (\u20130.20), "
        "non-smoker boost (+0.05), dependent boost (+0.10 per dependent for family products), and "
        "condition penalty (\u20130.05 per condition). Collaborative filtering boosts are applied based "
        "on profile-product affinities: postgraduate education with investment products (+0.12), low "
        "occupation hazard with investment products (+0.08), high expense ratio with protection products "
        "(+0.10), and priority-slider alignment (safety priority with protection, equity priority with "
        "investment)."
    )
    rb.p(
        "If a trained model is active, learned deltas from historical outcome data are incorporated using "
        "a weighted composition: product-level delta (55%), category delta (18%), subcategory delta (15%), "
        "and policy-type delta (12%). This enables the scoring to improve over time as more training data "
        "is provided."
    )
    rb.p("Task 3 \u2013 Coverage Prediction", bold=True)
    rb.p(
        "The predicted coverage amount is calculated based on the customer\u2019s stated protection purpose: "
        "Survivor Income uses 10\u00d7 annual expenses; Education Funding calculates years until the "
        "youngest child reaches 21 multiplied by annual expenses; Retirement Supplement uses 20\u00d7 "
        "annual expenses; and Estate Liquidity uses the maximum of 5\u00d7 annual expenses or 50% of net "
        "worth. All values are capped at LKR 50,000,000."
    )
    rb.p("Task 4 \u2013 Premium Calculation", bold=True)
    rb.p(
        "Monthly premium estimation uses a multi-factor model: base premium from the product, multiplied "
        "by a coverage factor (requested coverage / 1,000,000), adjusted by a risk multiplier (1.0 base + "
        "smoker surcharge +25%, age surcharge +15% if > 55, condition surcharge +5% each, occupation "
        "hazard surcharge). Rider premiums are added individually with the same risk multiplier. Tax is "
        "applied at the configured rate."
    )
    rb.p("Task 5 \u2013 Affordability and Lapse Prediction", bold=True)
    rb.p(
        "Affordability is scored on a 0.0\u20131.0 scale based on the premium-to-income ratio: 1.0 if "
        "\u22645%, 0.85 if \u226410%, 0.65 if \u226415%, 0.45 if \u226420%, 0.25 if \u226430%, and 0.10 "
        "otherwise. Lapse probability is estimated heuristically from a base of 0.10, increased by "
        "premium ratio, age, smoker status, and condition count factors."
    )
    rb.p("Task 6 \u2013 Policy Type Classification", bold=True)
    rb.p(
        "Each product is classified as Life, Retirement, Investment, or Critical Illness based on "
        "category/subcategory names, product tags, feature context (age, priority sliders), and "
        "product name patterns."
    )
    rb.p("Task 7 \u2013 Follow-Up Question Generation", bold=True)
    rb.p(
        "Adaptive follow-up questions are generated based on data completeness (missing required fields) "
        "and model confidence (score gap between top products below 0.12 threshold). Questions are "
        "contextual: if the top products are protection plans, questions about policy term and sum "
        "assured are generated; if education funding is selected but children\u2019s ages are missing, "
        "that question is prioritised."
    )
    rb.p(
        "The scoring pipeline is novel code, conceived and developed specifically for this project. "
        "The CART underwriting rules are adapted from standard insurance underwriting practices documented "
        "in the Life Insurance Association of Sri Lanka guidelines.",
        italic=True,
    )

    rb.h3("3. AiEngineClient \u2013 Backend-to-AI Integration")
    rb.p(
        "The AiEngineClient service in the Spring Boot backend handles all HTTP communication with the "
        "FastAPI AI engine. It uses Spring\u2019s RestClient for JSON requests (score, activate) and a "
        "raw HttpURLConnection fallback for multipart CSV uploads (train). The fallback was necessary "
        "because Spring\u2019s RestClient does not natively support multipart form-data uploads with "
        "file streams from disk. Error handling wraps all AI engine communication failures in "
        "RuntimeExceptions with descriptive messages, and logging at INFO level records all outbound "
        "calls for debugging."
    )
    rb.p("This integration pattern is novel code, developed for this project.", italic=True)

    rb.h3("4. Angular Lazy Loading and Route Guards")
    rb.p(
        "The Angular frontend uses standalone components with lazy-loaded routes configured in "
        "app.routes.ts. Each feature module is loaded on demand using dynamic import() expressions, "
        "reducing the initial bundle size. Three route guards enforce access control and journey "
        "sequencing: customerAuthGuard checks JWT validity, adminGuard verifies ADMIN role, and "
        "wizardProgressGuard enforces step completion order (step-1 \u2192 step-2 \u2192 step-3 \u2192 plan)."
    )
    rb.p(
        "The route guard implementation follows Angular\u2019s functional guard pattern (Angular 15+). "
        "This is adapted from Angular documentation with customisation for the wizard progress tracking.",
        italic=True,
    )

    rb.h3("5. Training Pipeline and Model Lifecycle")
    rb.p(
        "The training pipeline enables administrators to upload CSV files containing historical policy "
        "outcome data and derive updated scoring weights. The CSV must follow the policy-recommendation-v2 "
        "format with 12 required columns. Upon upload, the AI engine:\n"
        "\u2022 Parses and validates all rows, skipping those with missing required fields.\n"
        "\u2022 Groups rows by attributes: age buckets (young < 30, old > 55), smoker status, product code, "
        "category, subcategory, and policy type.\n"
        "\u2022 Calculates average outcome_score for each group.\n"
        "\u2022 Updates the DEFAULT_WEIGHTS dictionary based on group comparisons (e.g., if smokers have "
        "lower average outcomes than non-smokers, the smoker_penalty increases by the difference).\n"
        "\u2022 Generates per-product, per-category, and per-subcategory delta adjustments.\n"
        "\u2022 Persists the complete model artefact (weights + adjustments + metadata) as a JSON file.\n"
        "\u2022 Returns a training summary with rowsProcessed, updatedWeights, and modelArtifactId."
    )
    rb.p(
        "Model activation is a separate explicit step: the administrator promotes a trained model by "
        "calling POST /models/{artifactId}/activate, which loads the model\u2019s weights and adjustments "
        "into the AI engine\u2019s runtime memory. This separation ensures that training does not "
        "automatically affect production scoring."
    )
    rb.p("This training pipeline is novel code, designed for this project.", italic=True)

    rb.h3("6. Agentic Chat Subsystem (Open Chat)")
    rb.p(
        "The most significant new feature is the agentic AI chat subsystem, which provides a natural-language "
        "insurance advisory experience. The subsystem spans all four tiers of the architecture:"
    )
    rb.p("AI Engine Chat Module (ai-engine/app/chat/)", bold=True)
    rb.p(
        "The agentic chat loop is implemented in agent.py using a ReAct-style (Reasoning + Action) pattern. "
        "A Groq-hosted LLM (meta-llama/llama-4-scout-17b-16e-instruct) serves as the reasoning core, accessed "
        "via the OpenAI SDK. The agent operates in a loop: it receives the user message and conversation history, "
        "generates a response using the LLM, and if the LLM invokes a tool, the agent executes the tool, feeds "
        "the result back, and iterates until a final text response is produced."
    )
    rb.p("Five specialised tools are defined in tools.py:", bold=True)
    rb.bullet(
        "insurance_knowledge_base: Queries the MySQL database for product details, categories, eligibility rules, "
        "and pricing tables using parameterised SQL queries. Returns structured information about available products."
    )
    rb.bullet(
        "database_query: Executes read-only SQL queries against the insurance_university database with a "
        "100-row limit for safety. Provides general data access for answering data-driven questions."
    )
    rb.bullet(
        "web_search: Performs real-time web searches using the DuckDuckGo API to answer questions about "
        "external insurance topics, regulations, or market trends."
    )
    rb.bullet(
        "calculator: Evaluates mathematical expressions safely using the simpleeval library. Used for "
        "premium calculations, coverage estimates, and financial projections."
    )
    rb.bullet(
        "policy_scoring: Triggers the existing /score endpoint to generate AI-powered product recommendations "
        "based on a user's profile, bridging the agentic chat with the scoring pipeline."
    )
    rb.p("FAISS Vector Memory (memory.py)", bold=True)
    rb.p(
        "Long-term conversational memory is implemented using FAISS (Facebook AI Similarity Search) with "
        "sentence-transformers (all-MiniLM-L6-v2) for embedding generation. Each user-assistant exchange "
        "is embedded and stored in a per-user FAISS index. When a new conversation begins, relevant past "
        "interactions are retrieved via similarity search and injected into the system prompt as context. "
        "This enables the chat to maintain continuity across sessions without re-asking questions."
    )
    rb.p("Backend Chat Services", bold=True)
    rb.p(
        "The Spring Boot backend provides two chat modes through separate controllers and services:\n"
        "\u2022 OpenChatController: Handles the agentic open chat with endpoints for streaming (SSE), "
        "message history retrieval, session management, and memory clearing. It enforces per-user "
        "daily and monthly token limits.\n"
        "\u2022 CustomerChatService / ChatOrchestrator: Handles the wizard chat with two extraction modes "
        "(deterministic pattern matching and LLM-based extraction) configurable via app.chat.extractionMode. "
        "The DeterministicChatOrchestrator uses regex patterns to extract structured data from natural language, "
        "while the LlmChatOrchestrator delegates to the AI engine's /chat/message endpoint."
    )
    rb.p("Frontend Open Chat Component", bold=True)
    rb.p(
        "The Angular open-chat component provides a full-featured chat interface with SSE-based streaming. "
        "Messages are received token-by-token via EventSource and rendered in real-time with Markdown "
        "formatting. Tool calls are displayed with collapsible detail panels showing the tool name, "
        "arguments, and results. The component handles session management and message history."
    )
    rb.p(
        "The agentic chat subsystem is entirely novel code, designed and developed specifically for "
        "this project. It represents the most significant technical contribution beyond the original "
        "scoring pipeline.",
        italic=True,
    )

    rb.h3("7. New Skills Gained")
    rb.p(
        "The development of Insurance University required acquiring and applying several new technical "
        "skills beyond the standard curriculum:"
    )
    rb.bullet(
        "FastAPI Development: Building a production-grade Python REST API with Pydantic models, "
        "async endpoints, and file handling\u2014skills not covered in the BIS curriculum."
    )
    rb.bullet(
        "JWT Security Implementation: Designing a stateless authentication system with custom "
        "Spring Security filter chains, a skill that required deep understanding of the framework\u2019s "
        "security architecture."
    )
    rb.bullet(
        "Multi-Service Docker Compose Orchestration: Configuring health checks, volume mounts, and "
        "service dependencies across three containers."
    )
    rb.bullet(
        "AI/ML Scoring Pipeline Design: Implementing a multi-stage scoring algorithm that combines "
        "rule-based (CART) and data-driven (collaborative filtering) approaches."
    )
    rb.bullet(
        "Angular 21 Standalone Components: Adopting Angular\u2019s latest standalone component architecture "
        "and functional route guards, which differ significantly from the NgModule-based patterns."
    )
    rb.bullet(
        "Agentic AI Development: Implementing a ReAct-style tool-calling agent with LLM integration "
        "(Groq API via OpenAI SDK), custom tool definitions, and iterative reasoning loops."
    )
    rb.bullet(
        "Vector Databases and Embeddings: Building a FAISS-based semantic memory system with "
        "sentence-transformers for embedding generation and similarity-based retrieval."
    )
    rb.bullet(
        "Server-Sent Events (SSE): Implementing real-time streaming from backend to frontend for "
        "token-by-token chat response delivery."
    )

    # ── SECTION 14 ────────────────────────────────────────────────────────
    rb.add_page_break()
    rb.h1("14  Features to Be Tested / Not to Be Tested")
    rb.p("Table 7: Test Plan \u2013 Features to Be Tested", bold=True, italic=True)
    rb.table(
        ["Category", "Feature", "Test Type", "Priority"],
        [
            ["Authentication", "Customer registration with valid data", "Functional", "High"],
            ["Authentication", "Customer login with valid credentials", "Functional", "High"],
            ["Authentication", "Admin login with valid credentials", "Functional", "High"],
            ["Authentication", "Access protected endpoint without JWT", "Security", "High"],
            ["Authentication", "Access admin endpoint with USER role JWT", "Security", "High"],
            ["Wizard", "Step 1 form validation and answer persistence", "Functional", "High"],
            ["Wizard", "Step 2 guard enforcement (requires Step 1)", "Functional", "Medium"],
            ["Wizard", "Step 3 answer persistence and recommendation trigger", "Functional", "High"],
            ["Recommendations", "AI scoring with valid input (5 products)", "Functional/Integration", "High"],
            ["Recommendations", "Eligibility rule filtering", "Functional", "High"],
            ["Recommendations", "Recommendation audit trail persistence", "Functional", "Medium"],
            ["Documents", "KYC document upload (NIC front/back, medical)", "Functional", "High"],
            ["Documents", "Document version management", "Functional", "Medium"],
            ["Documents", "Document download access control", "Security", "High"],
            ["Admin", "Product CRUD operations", "Functional", "High"],
            ["Admin", "Eligibility rule CRUD operations", "Functional", "High"],
            ["Admin", "Training dataset upload and model creation", "Integration", "High"],
            ["Admin", "Model promotion (activate)", "Functional", "Medium"],
            ["Admin", "Log search and export", "Functional", "Medium"],
            ["AI Engine", "/score endpoint contract validation", "Contract/Integration", "High"],
            ["AI Engine", "/train endpoint with valid CSV", "Contract/Integration", "High"],
            ["AI Engine", "CART underwriting rule correctness", "Unit", "High"],
            ["Performance", "Recommendation response time (< 5s for 5 products)", "Performance", "Medium"],
            ["Open Chat", "Agentic SSE streaming with tool calls", "Functional/Integration", "High"],
            ["Open Chat", "FAISS memory retrieval across sessions", "Functional", "Medium"],
            ["Open Chat", "Token limit enforcement (daily/monthly)", "Functional", "Medium"],
            ["Open Chat", "Tool invocation: knowledge base, calculator, web search", "Integration", "High"],
            ["Wizard Chat", "Deterministic extraction of structured answers", "Functional", "High"],
            ["Wizard Chat", "LLM-based extraction mode fallback", "Integration", "Medium"],
        ],
    )
    rb.p("")
    rb.p("Features Not to Be Tested", bold=True)
    rb.table(
        ["Feature", "Reason"],
        [
            ["Payment gateway integration", "Not implemented in current prototype"],
            ["IRCSL regulatory API integration", "Not implemented; future enhancement"],
            ["Email/SMS notification delivery", "Not implemented; future enhancement"],
            ["NIC OCR verification", "Not implemented; future enhancement"],
            ["Load testing (> 100 concurrent users)", "Out of scope for prototype evaluation"],
            ["Cross-browser compatibility (Safari, Firefox)", "Limited to Chrome for prototype"],
            ["Mobile native app", "Only responsive web app; no native app built"],
        ],
    )

    # ── SECTION 15 ────────────────────────────────────────────────────────
    rb.add_page_break()
    rb.h1("15  Test Cases")
    rb.p("Table 8: Test Cases", bold=True, italic=True)

    test_cases = [
        {
            "id": "TC001",
            "desc": "Verify customer registration with valid credentials",
            "pre": "User is on the registration page. No account exists for the test email.",
            "steps": "a. Enter name: 'Test User'\nb. Enter email: 'test@example.com'\nc. Enter password: 'Test@1234'\nd. Click Register button.",
            "expected": "User is registered, JWT token returned, user redirected to wizard step 1.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if HTTP 200 returned with valid JWT and user record created in database.",
            "data": "Name: Test User; Email: test@example.com; Password: Test@1234",
            "post": "User account exists in database with BCrypt-hashed password.",
            "notes": "Ensure email uniqueness constraint is tested in TC002."
        },
        {
            "id": "TC002",
            "desc": "Verify duplicate email registration is rejected",
            "pre": "User 'test@example.com' already exists in the database.",
            "steps": "a. Enter name: 'Another User'\nb. Enter email: 'test@example.com'\nc. Enter password: 'Another@1234'\nd. Click Register button.",
            "expected": "Registration fails with HTTP 409 Conflict error message.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if error response returned and no duplicate record created.",
            "data": "Name: Another User; Email: test@example.com; Password: Another@1234",
            "post": "Only one user record exists for test@example.com.",
            "notes": "Tests database unique constraint on email field."
        },
        {
            "id": "TC003",
            "desc": "Verify customer login with valid credentials",
            "pre": "User 'test@example.com' is registered.",
            "steps": "a. Enter email: 'test@example.com'\nb. Enter password: 'Test@1234'\nc. Click Login button.",
            "expected": "User is logged in, JWT token returned, user redirected to dashboard.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if HTTP 200 returned with valid JWT containing USER role.",
            "data": "Email: test@example.com; Password: Test@1234",
            "post": "JWT stored in localStorage for subsequent API calls.",
            "notes": "Verify JWT payload contains correct email and role claims."
        },
        {
            "id": "TC004",
            "desc": "Verify access to protected endpoint without JWT",
            "pre": "No JWT token in request headers.",
            "steps": "a. Send GET /api/customer/sessions without Authorization header.",
            "expected": "HTTP 401 Unauthorized returned.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if 401 returned and no session data exposed.",
            "data": "No authentication data.",
            "post": "No session created or data accessed.",
            "notes": "Critical security test."
        },
        {
            "id": "TC005",
            "desc": "Verify wizard Step 1 answer persistence",
            "pre": "Customer is logged in with active session.",
            "steps": "a. Navigate to /wizard/step-1\nb. Enter age: 35\nc. Set smoker: No\nd. Select conditions: None\ne. Set dependents: 2\nf. Click Next.",
            "expected": "Answers saved via POST /api/customer/sessions/{id}/answers. User redirected to Step 2.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if answers retrievable via GET and values match input.",
            "data": "age: 35, smoker: false, conditions: [], dependents: 2",
            "post": "CustomerAnswerEntity records created for each key-value pair.",
            "notes": "Verify JSON serialisation of complex values."
        },
        {
            "id": "TC006",
            "desc": "Verify wizard Step 2 guard blocks access without Step 1 completion",
            "pre": "Customer is logged in. Step 1 NOT completed.",
            "steps": "a. Navigate directly to /wizard/step-2 via URL.",
            "expected": "User is redirected back to /wizard/step-1 by wizardProgressGuard.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if user cannot access Step 2 and is redirected.",
            "data": "Direct URL navigation to /wizard/step-2",
            "post": "User remains on Step 1.",
            "notes": "Tests Angular route guard functionality."
        },
        {
            "id": "TC007",
            "desc": "Verify AI recommendation generation with valid inputs",
            "pre": "Customer has completed all 3 wizard steps. All answers saved. AI engine running.",
            "steps": "a. Trigger POST /api/customer/sessions/{id}/recommendations.\nb. Wait for response.",
            "expected": "HTTP 200 with rankedProducts array (1-5 products), each with score, premium, affordability, reasons.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if products ranked by score descending, all fields present, response < 5 seconds.",
            "data": "Full wizard answers: age=35, smoker=false, income=100000, expenses=50000, netWorth=500000, dependents=2, purpose=SurvivorIncome",
            "post": "RecommendationRunEntity persisted. SessionLogEntity created with RECOMMENDATIONS_FETCHED event.",
            "notes": "Integration test spanning backend and AI engine."
        },
        {
            "id": "TC008",
            "desc": "Verify KYC document upload",
            "pre": "Customer logged in with active session.",
            "steps": "a. Navigate to /proposal/upload\nb. Select NIC front image file (JPEG, < 5MB)\nc. Upload via POST /api/customer/sessions/{id}/documents?docType=nic&docSide=front.",
            "expected": "HTTP 200 with document metadata. File stored on filesystem.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if document record created, file exists at stored path, latestForSession=true.",
            "data": "File: test_nic_front.jpg (200KB JPEG); docType: nic; docSide: front",
            "post": "CustomerDocumentEntity created with version 1.",
            "notes": "Test re-upload to verify version increment."
        },
        {
            "id": "TC009",
            "desc": "Verify admin training dataset upload and model creation",
            "pre": "Admin logged in. AI engine running. Valid CSV file available.",
            "steps": "a. Navigate to Admin > Training\nb. Upload CSV file with 100 rows in policy-recommendation-v2 format\nc. Wait for training completion.",
            "expected": "HTTP 200 with training results: rowsProcessed=100, modelArtifactId, updatedWeights.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if DatasetMetaEntity and ModelVersionEntity created. AI engine returns valid artifact ID.",
            "data": "CSV with 100 rows, 12 required columns, valid data types.",
            "post": "Model artifact JSON file created in AI engine data/models/ directory.",
            "notes": "Integration test spanning backend and AI engine. Test with invalid CSV in TC010."
        },
        {
            "id": "TC010",
            "desc": "Verify admin cannot access customer endpoints",
            "pre": "Admin logged in with ADMIN role JWT.",
            "steps": "a. Send POST /api/customer/sessions with admin JWT.",
            "expected": "Request should be processed based on security config (customer endpoints check authentication, not role restriction for session creation).",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if system correctly handles cross-role access per security configuration.",
            "data": "Admin JWT token",
            "post": "Verify expected behaviour matches security filter chain rules.",
            "notes": "Documents actual RBAC behaviour for cross-role scenarios."
        },
        {
            "id": "TC011",
            "desc": "Verify agentic open chat with SSE streaming and tool invocation",
            "pre": "Customer logged in. AI engine running with Groq API key configured.",
            "steps": "a. Navigate to /open-chat\nb. Type: 'What life insurance products are available?'\nc. Send message\nd. Observe SSE stream response.",
            "expected": "SSE stream delivers tokens incrementally. Agent invokes insurance_knowledge_base tool. Final response lists available products with details.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if SSE stream completes, tool call visible in UI, response contains accurate product data from database.",
            "data": "User message: 'What life insurance products are available?'",
            "post": "OpenChatMessage records created (USER + ASSISTANT + TOOL roles). Token usage tracked.",
            "notes": "Integration test spanning frontend SSE, backend streaming, AI engine agent loop, and Groq API."
        },
        {
            "id": "TC012",
            "desc": "Verify wizard chat deterministic extraction mode",
            "pre": "Customer logged in. Session created. app.chat.extractionMode=deterministic.",
            "steps": "a. Navigate to /chat\nb. Type: 'I am 32 years old, non-smoker, married with 2 kids'\nc. Send message.",
            "expected": "System extracts structured answers: age=32, smoker=false, dependents=2. Extracted fields saved as CustomerAnswerEntity records.",
            "actual": "[To be filled during testing]",
            "criteria": "Pass if deterministic parser correctly identifies age, smoker status, and dependents from natural language.",
            "data": "User message: 'I am 32 years old, non-smoker, married with 2 kids'",
            "post": "CustomerAnswerEntity records created for extracted key-value pairs.",
            "notes": "Tests DeterministicChatOrchestrator regex patterns."
        },
    ]

    for tc in test_cases:
        rb.p(f"Test Case {tc['id']}", bold=True)
        rb.table(
            ["Element", "Detail"],
            [
                ["Test Case ID", tc["id"]],
                ["Test Description", tc["desc"]],
                ["Preconditions", tc["pre"]],
                ["Test Steps", tc["steps"]],
                ["Expected Result", tc["expected"]],
                ["Actual Result", tc["actual"]],
                ["Pass/Fail Criteria", tc["criteria"]],
                ["Test Data", tc["data"]],
                ["Post-Conditions", tc["post"]],
                ["Notes/Comments", tc["notes"]],
            ],
        )
        rb.p("")

    # ═══════════════════════════════════════════════════════════════════════
    # PART V: PROJECT SUMMARY
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("V  Project Summary")

    # ── SECTION 16 ────────────────────────────────────────────────────────
    rb.h1("16  Delivering Sustainability")
    rb.p(
        "Insurance University contributes to sustainability across environmental, social, and economic "
        "dimensions, aligning with several United Nations Sustainable Development Goals (SDGs)."
    )
    rb.p("Environmental Sustainability", bold=True)
    rb.p(
        "By digitising the insurance advisory and proposal submission process, Insurance University "
        "eliminates the need for physical paper forms, printed product brochures, and hard-copy KYC "
        "document submissions. A typical insurance application in Sri Lanka involves 5\u201310 pages of "
        "paper forms and 3\u20135 photocopied identity/medical documents. For an insurance company "
        "processing 10,000 applications annually, this digital alternative could prevent the consumption "
        "of approximately 80,000\u2013150,000 sheets of paper per year. Additionally, the platform reduces "
        "the need for customers to travel to insurance offices or agent locations, thereby lowering "
        "carbon emissions associated with transportation."
    )
    rb.p("Social Sustainability", bold=True)
    rb.p(
        "The platform promotes financial inclusion by making insurance advisory services accessible to "
        "anyone with an internet connection, regardless of geographic location. In Sri Lanka, where "
        "insurance agents are concentrated in urban areas, rural customers face significant barriers to "
        "accessing professional insurance advice. Insurance University\u2019s web-based wizard and AI "
        "scoring engine provide the same quality of assessment to all users, irrespective of location. "
        "The explainable AI approach also empowers customers with transparent, understandable "
        "information, fostering financial literacy in a historically opaque industry."
    )
    rb.p("Economic Sustainability", bold=True)
    rb.p(
        "For insurance companies, the platform offers a scalable, low-cost distribution channel that "
        "reduces customer acquisition costs. The AI-driven recommendation engine can serve thousands "
        "of concurrent users without proportional increases in human advisory staff. The admin console\u2019s "
        "training and model management capabilities enable continuous improvement of recommendation "
        "accuracy, potentially reducing policy lapse rates and improving customer lifetime value. "
        "The unmatched needs analytics feature provides data-driven insights for product development, "
        "enabling insurers to identify and address gaps in their product portfolios."
    )
    rb.p("Alignment with UN SDGs", bold=True)
    rb.bullet(
        "SDG 3 (Good Health and Well-being): By recommending appropriate medical and life insurance "
        "products, the platform supports access to health protection mechanisms."
    )
    rb.bullet(
        "SDG 8 (Decent Work and Economic Growth): The platform contributes to financial inclusion "
        "and efficient insurance distribution, supporting economic resilience."
    )
    rb.bullet(
        "SDG 9 (Industry, Innovation, and Infrastructure): The AI-powered recommendation engine "
        "represents innovation in the insurance technology (InsurTech) domain."
    )
    rb.bullet(
        "SDG 12 (Responsible Consumption and Production): The paperless process reduces resource "
        "consumption in the insurance value chain."
    )

    # ── SECTION 17 ────────────────────────────────────────────────────────
    rb.h1("17  Conclusions and Open Issues")
    rb.p("Achievements", bold=True)
    rb.p(
        "Insurance University successfully delivers a functional, end-to-end AI-powered insurance "
        "recommendation platform that addresses the core problems identified in the Sri Lankan insurance "
        "advisory process. The key achievements of the project are:"
    )
    rb.numbered(
        "Four-Tier Microservices Architecture: A production-ready architecture comprising an Angular 21 "
        "SPA, Spring Boot 3.3.5 REST API, FastAPI AI engine with an agentic chat subsystem, and MySQL 8.4, "
        "orchestrated via Docker Compose. This architecture supports independent scaling and deployment of each tier."
    )
    rb.numbered(
        "Multi-Stage AI Scoring Engine: A comprehensive scoring pipeline that combines CART-based "
        "underwriting rules with collaborative filtering heuristics, trainable feature weights, coverage "
        "prediction, premium estimation, affordability analysis, and lapse probability estimation. The "
        "engine produces explainable, transparent recommendations with human-readable reasons."
    )
    rb.numbered(
        "Agentic AI Chat Subsystem: A ReAct-style conversational AI agent powered by Groq (LLaMA 4 Scout) "
        "with five specialised tools (insurance knowledge base, database query, web search, calculator, "
        "policy scoring), FAISS-based long-term vector memory using sentence-transformers, and real-time "
        "SSE streaming. This subsystem enables natural-language insurance advisory beyond the structured wizard."
    )
    rb.numbered(
        "Trainable Model Architecture: An end-to-end model lifecycle management system that enables "
        "administrators to upload historical policy outcome data, train updated scoring weights, and "
        "promote models to production\u2014all through the web-based admin console."
    )
    rb.numbered(
        "Dual Chat Modes: The platform offers two conversational interfaces\u2014a Wizard Chat with "
        "deterministic and LLM-based data extraction for structured answer capture, and an Open Chat "
        "with full agentic capabilities for free-form insurance advisory."
    )
    rb.numbered(
        "Comprehensive Customer Journey: A structured wizard interface (3 steps) with progress guards, "
        "followed by ranked recommendations, side-by-side product comparison, premium simulation, "
        "KYC document upload with version management, and proposal summary\u2014covering the complete "
        "insurance discovery-to-application flow."
    )
    rb.numbered(
        "Robust Security Implementation: Stateless JWT authentication with RBAC (Customer/Admin roles), "
        "BCrypt password hashing, SHA-256 session log anonymisation, document access control, and "
        "sensitive field redaction in logs."
    )
    rb.numbered(
        "Administrative Console: A feature-rich back-office interface for product catalogue management, "
        "eligibility rule configuration, training dataset and model management, searchable/exportable "
        "event logs, and unmatched customer needs tracking."
    )
    rb.numbered(
        "Full Audit Trail: Every AI scoring call is persisted with complete request and response "
        "payloads, and all customer journey events are logged with anonymised identifiers, supporting "
        "regulatory compliance and post-hoc analysis."
    )
    rb.p("Open Issues and Future Work", bold=True)
    rb.p(
        "Despite the achievements outlined above, several issues remain open and represent opportunities "
        "for future enhancement:"
    )
    rb.numbered(
        "Advanced RAG Integration: While the current FAISS memory provides conversational continuity, "
        "a full Retrieval-Augmented Generation (RAG) pipeline with insurance policy document embeddings "
        "would enable the agentic chat to cite specific policy terms, conditions, and exclusion clauses "
        "directly from product documentation."
    )
    rb.numbered(
        "Payment Gateway Integration: The current prototype does not include online premium payment "
        "capabilities. Integration with Sri Lankan payment gateways (PayHere, LankaPay) would complete "
        "the end-to-end purchase journey."
    )
    rb.numbered(
        "IRCSL API Connectivity: The platform does not currently connect to IRCSL\u2019s regulatory systems "
        "for product registration validation or compliance reporting. This integration would be "
        "essential for production deployment."
    )
    rb.numbered(
        "Limited Test Coverage: The current prototype has minimal automated test coverage (both backend "
        "unit tests and frontend component tests). A comprehensive test suite should be developed before "
        "production deployment, including integration tests for the backend-AI engine communication."
    )
    rb.numbered(
        "Model Persistence Across Restarts: The AI engine\u2019s active model weights are held in memory "
        "and loaded from disk on startup. If the engine restarts without explicit model activation, "
        "it reverts to default weights. An auto-load mechanism for the last active model should be "
        "implemented."
    )
    rb.numbered(
        "Multi-Tenant Support: The current system supports a single insurance company (single admin "
        "account). Multi-tenant architecture would enable multiple insurers to manage their own product "
        "catalogues and training data independently."
    )
    rb.numbered(
        "Advanced Analytics Dashboard: The admin dashboard currently provides basic KPIs. Future "
        "versions should include conversion funnel analysis, recommendation acceptance rates, "
        "product popularity trends, and customer demographic insights."
    )
    rb.numbered(
        "Production Deployment Hardening: The current configuration uses development defaults "
        "(plaintext JWT secrets, in-memory uploads, localhost CORS origins). Production deployment "
        "requires HTTPS enforcement, secret management (e.g., HashiCorp Vault), cloud storage for "
        "documents (e.g., AWS S3), and robust error monitoring (e.g., Sentry)."
    )

    # ═══════════════════════════════════════════════════════════════════════
    # GLOSSARY
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("VI  Glossary")
    rb.p("Table 9: Glossary of Terms", bold=True, italic=True)
    rb.table(
        ["Term", "Definition"],
        [
            ["Agentic AI", "An AI system that autonomously selects and executes tools to fulfil user requests, using iterative reasoning loops."],
            ["API", "Application Programming Interface \u2013 a set of rules that allows software components to communicate."],
            ["CART", "Classification and Regression Tree \u2013 a decision tree algorithm used for classification and prediction."],
            ["CORS", "Cross-Origin Resource Sharing \u2013 a security mechanism that controls which origins can access a web API."],
            ["CRUD", "Create, Read, Update, Delete \u2013 the four basic operations on persistent data."],
            ["CSRF", "Cross-Site Request Forgery \u2013 an attack that tricks a user into submitting unintended requests."],
            ["Docker", "A platform for building, running, and managing containerised applications."],
            ["Docker Compose", "A tool for defining and running multi-container Docker applications using a YAML configuration file."],
            ["FAISS", "Facebook AI Similarity Search \u2013 a library for efficient similarity search and clustering of dense vectors."],
            ["FastAPI", "A modern Python web framework for building APIs with automatic validation and documentation."],
            ["Groq", "A cloud inference platform providing high-speed access to large language models via an OpenAI-compatible API."],
            ["HMAC-SHA256", "Hash-based Message Authentication Code using SHA-256 \u2013 a cryptographic algorithm for signing JWTs."],
            ["IRCSL", "Insurance Regulatory Commission of Sri Lanka \u2013 the statutory body regulating insurance in Sri Lanka."],
            ["JPA", "Java Persistence API \u2013 a specification for object-relational mapping in Java applications."],
            ["JWT", "JSON Web Token \u2013 a compact, URL-safe token for securely transmitting claims between parties."],
            ["KYC", "Know Your Customer \u2013 the process of verifying customer identity and assessing risk."],
            ["LKR", "Sri Lankan Rupee \u2013 the official currency of Sri Lanka."],
            ["PDPA", "Personal Data Protection Act \u2013 Sri Lanka\u2019s data protection legislation."],
            ["PEP", "Politically Exposed Person \u2013 an individual in a prominent public role, subject to enhanced due diligence."],
            ["RBAC", "Role-Based Access Control \u2013 a method of restricting system access based on user roles."],
            ["ReAct", "Reasoning and Action \u2013 a prompting paradigm where an LLM alternates between reasoning steps and tool actions."],
            ["REST", "Representational State Transfer \u2013 an architectural style for designing networked applications."],
            ["SPA", "Single Page Application \u2013 a web application that loads a single HTML page and dynamically updates content."],
            ["Spring Boot", "A Java framework that simplifies the development of production-ready Spring-based applications."],
            ["SSE", "Server-Sent Events \u2013 a standard for pushing real-time updates from server to client over HTTP."],
            ["UUID", "Universally Unique Identifier \u2013 a 128-bit label used for unique identification."],
            ["Vector Embedding", "A dense numerical representation of text that captures semantic meaning for similarity search."],
            ["WCAG", "Web Content Accessibility Guidelines \u2013 international standards for web accessibility."],
        ],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("VII  References / Bibliography")
    refs = [
        "[1]\tAlexander, I. and Robertson, S. (2004) 'Understanding Project Sociology by Modeling Stakeholders', IEEE Software, 21(1), pp. 23\u201327.",
        "[2]\tAngular Team (2026) Angular Documentation. Available at: https://angular.dev (Accessed: 15 April 2026).",
        "[3]\tBasel Committee on Banking Supervision (2017) 'Customer DueLigence for Banks', Bank for International Settlements.",
        "[4]\tBreiman, L. et al. (1984) Classification and Regression Trees. Boca Raton: CRC Press.",
        "[5]\tFowler, M. (2004) UML Distilled. 3rd edn. Boston: Pearson Education.",
        "[6]\tGovernment of Sri Lanka (2000) Regulation of Insurance Industry Act No. 43 of 2000. Colombo: Government Press.",
        "[7]\tGovernment of Sri Lanka (2022) Personal Data Protection Act No. 9 of 2022. Colombo: Government Press.",
        "[8]\tJohnson, J., Douze, M. and J\u00e9gou, H. (2021) 'Billion-Scale Similarity Search with GPUs', IEEE Transactions on Big Data, 7(3), pp. 535\u2013547.",
        "[9]\tJones, C. and Sheridan, M. (2015) 'Insurance Underwriting Decision Models: A Survey', Journal of Risk and Insurance, 82(3), pp. 621\u2013645.",
        "[10]\tOWASP Foundation (2021) OWASP Top Ten. Available at: https://owasp.org/www-project-top-ten/ (Accessed: 10 April 2026).",
        "[11]\tRamakrishnan, R. and Gehrke, J. (2003) Database Management Systems. 3rd edn. New York: McGraw-Hill.",
        "[12]\tRichardson, C. (2018) Microservices Patterns. Shelter Island: Manning Publications.",
        "[13]\tRobertson, S. and Robertson, J. (2012) Mastering the Requirements Process. 3rd edn. Upper Saddle River: Addison-Wesley.",
        "[14]\tSommerville, I. (2016) Software Engineering. 10th edn. Harlow: Pearson Education.",
        "[15]\tSpring Team (2024) Spring Boot Reference Documentation. Available at: https://docs.spring.io/spring-boot/docs/current/reference/html/ (Accessed: 12 April 2026).",
        "[16]\tSwiss Re Institute (2023) 'World Insurance: The Recovery Gains Pace', Sigma No. 3/2023.",
        "[17]\tTiram\u00edsou Sebastien (2023) FastAPI Documentation. Available at: https://fastapi.tiangolo.com (Accessed: 12 April 2026).",
        "[18]\tW3C (2018) Web Content Accessibility Guidelines (WCAG) 2.1. Available at: https://www.w3.org/TR/WCAG21/ (Accessed: 14 April 2026).",
        "[19]\tYao, S. et al. (2023) 'ReAct: Synergizing Reasoning and Acting in Language Models', ICLR 2023.",
        "[20]\tGroq Inc. (2025) Groq API Documentation. Available at: https://console.groq.com/docs (Accessed: 15 April 2026).",
    ]
    for ref in refs:
        rb.p(ref)

    # ═══════════════════════════════════════════════════════════════════════
    # APPENDICES
    # ═══════════════════════════════════════════════════════════════════════
    rb.add_page_break()
    rb.h1("Appendices")

    rb.h2("Appendix A: AI Scoring Weight Parameters")
    rb.p("Table 5: AI Scoring Weight Parameters (Default Values)", bold=True, italic=True)
    rb.table(
        ["Weight Parameter", "Default Value", "Effect", "Description"],
        [
            ["age_young_boost", "0.10", "+score", "Boost for applicants under 30 years"],
            ["age_old_penalty", "0.15", "\u2013score", "Penalty for applicants over 55 years"],
            ["smoker_penalty", "0.20", "\u2013score", "Penalty for tobacco users"],
            ["nonsmoker_boost", "0.05", "+score", "Discount for non-tobacco users"],
            ["dependent_boost", "0.10/dep", "+score", "Per-dependent boost for family products"],
            ["ratio_high_penalty", "0.10", "\u2013score", "Penalty when coverage > 10\u00d7 annual income"],
            ["ratio_low_boost", "0.05", "+score", "Boost for moderate coverage ratios"],
            ["senior_cap_penalty", "0.30", "\u2013score", "Penalty for age > 65 on non-senior products"],
            ["condition_penalty", "0.05/cond", "\u2013score", "Per-condition health risk penalty"],
            ["smoker_risk_pct", "0.25", "+premium", "25% premium surcharge for smokers"],
            ["age_old_risk_pct", "0.15", "+premium", "15% premium surcharge for age > 55"],
            ["condition_risk_pct", "0.05/cond", "+premium", "5% premium surcharge per condition"],
            ["postgrad_invest_boost", "0.12", "+score", "Postgrad + investment product affinity"],
            ["undergrad_invest_boost", "0.06", "+score", "Undergrad + investment product affinity"],
            ["low_hazard_invest_boost", "0.08", "+score", "Low-risk occupation + investment affinity"],
            ["high_expense_protection", "0.10", "+score", "High expense ratio + protection affinity"],
            ["low_networth_protection", "0.08", "+score", "Low net worth + protection affinity"],
            ["afford_penalty", "0.15", "\u2013score", "Penalty when premium > 20% disposable income"],
            ["safety_priority_boost", "0.04", "+score", "Safety priority \u2265 4 + protection affinity"],
            ["equity_priority_invest", "0.05", "+score", "Equity priority \u2265 4 + investment affinity"],
            ["family_history_penalty", "0.05", "\u2013score", "Family medical history penalty"],
        ],
    )

    rb.h2("Appendix B: API Endpoint Summary")
    rb.p("Complete list of REST API endpoints exposed by the Spring Boot backend:", bold=True)
    rb.table(
        ["Method", "Endpoint", "Auth", "Description"],
        [
            ["POST", "/api/auth/register", "Public", "Customer registration"],
            ["POST", "/api/auth/login", "Public", "Customer login"],
            ["POST", "/api/auth/admin/login", "Public", "Admin login"],
            ["GET", "/api/auth/me", "JWT", "Get authenticated user profile"],
            ["POST", "/api/customer/sessions", "JWT (USER)", "Create new session"],
            ["GET", "/api/customer/sessions", "JWT (USER)", "List user sessions"],
            ["GET", "/api/customer/sessions/{id}", "JWT (USER)", "Get session details"],
            ["DELETE", "/api/customer/sessions/{id}", "JWT (USER)", "Delete session"],
            ["PATCH", "/api/customer/sessions/{id}/complete", "JWT (USER)", "Mark session completed"],
            ["POST", "/api/customer/sessions/{id}/answers", "JWT (USER)", "Save wizard answers"],
            ["GET", "/api/customer/sessions/{id}/answers", "JWT (USER)", "Get answers as map"],
            ["POST", "/api/customer/sessions/{id}/recommendations", "JWT (USER)", "Trigger AI scoring"],
            ["GET", "/api/customer/sessions/{id}/recommendations/latest", "JWT (USER)", "Get last recommendation"],
            ["POST", "/api/customer/sessions/{id}/documents", "JWT (USER)", "Upload KYC document"],
            ["GET", "/api/customer/sessions/{id}/documents", "JWT (USER)", "List session documents"],
            ["GET", "/api/customer/sessions/{id}/documents/{docId}/download", "JWT (USER)", "Download document"],
            ["GET", "/api/customer/documents/latest", "JWT (USER)", "Get reusable user documents"],
            ["POST", "/api/customer/sessions/{id}/chat", "JWT (USER)", "Send chat message"],
            ["GET", "/api/customer/sessions/{id}/chat/history", "JWT (USER)", "Get chat history"],
            ["POST", "/api/customer/feedback", "JWT (USER)", "Submit feedback"],
            ["POST", "/api/logs", "Public", "Client event logging"],
            ["GET", "/api/admin/products", "JWT (ADMIN)", "List products"],
            ["POST", "/api/admin/products", "JWT (ADMIN)", "Create product"],
            ["GET", "/api/admin/products/{id}", "JWT (ADMIN)", "Get product"],
            ["PUT", "/api/admin/products/{id}", "JWT (ADMIN)", "Update product"],
            ["DELETE", "/api/admin/products/{id}", "JWT (ADMIN)", "Delete product"],
            ["GET", "/api/admin/rules", "JWT (ADMIN)", "List eligibility rules"],
            ["POST", "/api/admin/rules", "JWT (ADMIN)", "Create rule"],
            ["GET", "/api/admin/rules/{id}", "JWT (ADMIN)", "Get rule"],
            ["PUT", "/api/admin/rules/{id}", "JWT (ADMIN)", "Update rule"],
            ["DELETE", "/api/admin/rules/{id}", "JWT (ADMIN)", "Delete rule"],
            ["POST", "/api/admin/training/datasets", "JWT (ADMIN)", "Upload training CSV"],
            ["GET", "/api/admin/training/datasets", "JWT (ADMIN)", "List datasets"],
            ["POST", "/api/admin/training/datasets/{id}/retrain", "JWT (ADMIN)", "Retrain from dataset"],
            ["POST", "/api/admin/training/models", "JWT (ADMIN)", "Create model record"],
            ["GET", "/api/admin/training/models", "JWT (ADMIN)", "List models"],
            ["POST", "/api/admin/training/models/{id}/promote", "JWT (ADMIN)", "Activate model"],
            ["GET", "/api/admin/logs", "JWT (ADMIN)", "Search logs"],
            ["GET", "/api/admin/logs/export", "JWT (ADMIN)", "Export logs (CSV/JSON)"],
            ["GET", "/api/admin/insights/unmatched-needs", "JWT (ADMIN)", "List unmatched needs"],
            ["POST", "/api/admin/insights/unmatched-needs", "JWT (ADMIN)", "Record unmatched need"],
            ["DELETE", "/api/admin/insights/unmatched-needs/{id}", "JWT (ADMIN)", "Delete unmatched need"],
            ["GET", "/api/admin/pricing-tables", "JWT (ADMIN)", "List pricing tables"],
            ["POST", "/api/admin/pricing-tables", "JWT (ADMIN)", "Create pricing table"],
            ["PUT", "/api/admin/pricing-tables/{id}", "JWT (ADMIN)", "Update pricing table"],
            ["DELETE", "/api/admin/pricing-tables/{id}", "JWT (ADMIN)", "Delete pricing table"],
            ["POST", "/api/admin/dev/seed", "JWT (ADMIN)", "Seed default products"],
            ["GET", "/api/open-chat/stream", "JWT (USER)", "SSE stream agentic chat response"],
            ["GET", "/api/open-chat/messages", "JWT (USER)", "Get open chat message history"],
            ["GET", "/api/open-chat/sessions", "JWT (USER)", "List open chat sessions"],
            ["DELETE", "/api/open-chat/sessions/{id}", "JWT (USER)", "Delete open chat session"],
            ["DELETE", "/api/open-chat/memory", "JWT (USER)", "Clear FAISS vector memory"],
        ],
    )

    rb.h2("Appendix C: AI Engine Endpoints")
    rb.table(
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/health", "Health check with active model metadata"],
            ["POST", "/score", "Score products for a customer session"],
            ["POST", "/train", "Train model from CSV file upload"],
            ["POST", "/models/{artifactId}/activate", "Promote a trained model to active"],
            ["POST", "/chat/stream", "SSE-streamed agentic chat with tool calling"],
            ["POST", "/chat/message", "Non-streaming chat message (wizard LLM extraction)"],
            ["GET", "/chat/health", "Chat subsystem health check (Groq/FAISS status)"],
            ["DELETE", "/chat/memory/{user_id}", "Clear user\u2019s FAISS vector memory"],
        ],
    )

    rb.h2("Appendix D: Non-Functional Requirements")
    rb.p("Table 3: Non-Functional Requirements", bold=True, italic=True)
    rb.table(
        ["ID", "Requirement", "Category", "Target"],
        [
            ["NFR-01", "Recommendation response time", "Performance", "< 5 seconds for 5 products"],
            ["NFR-02", "Frontend initial load time", "Performance", "< 3 seconds on broadband"],
            ["NFR-03", "System availability", "Availability", "99% uptime (excluding planned maintenance)"],
            ["NFR-04", "Password storage", "Security", "BCrypt with default strength (10 rounds)"],
            ["NFR-05", "Token expiration", "Security", "24-hour JWT expiry"],
            ["NFR-06", "Session log anonymisation", "Privacy", "SHA-256 hashing of session identifiers"],
            ["NFR-07", "Responsive design", "Usability", "Functional on 320px\u20131920px viewports"],
            ["NFR-08", "Accessibility", "Usability", "WCAG 2.1 Level AA compliance"],
            ["NFR-09", "Data backup", "Reliability", "MySQL volume persistence via Docker"],
            ["NFR-10", "Browser support", "Compatibility", "Chrome 90+, Edge 90+ (primary)"],
        ],
    )

    # ═══════════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════════
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), OUTPUT_FILE)
    rb.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_report()
    print(f"\nDone! Report generated at:\n  {path}")
    print("\nPost-generation steps:")
    print("  1. Open the .docx in Microsoft Word")
    print("  2. Right-click Table of Contents \u2192 Update Field \u2192 Update Entire Table")
    print("  3. Insert diagrams/screenshots at [INSERT IMAGE: ...] placeholders")
    print("  4. Update word count in Declaration section")
    print("  5. Review and adjust page count")
